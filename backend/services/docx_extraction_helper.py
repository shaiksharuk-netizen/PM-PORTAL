"""
Helper functions for extracting text from DOCX files
Simplified version: Only extracts plain text and applies hardcoded playbook links mapping.
"""
import io
from docx import Document

# Import playbook links mapping
try:
    from services.playbook_links_mapping import enrich_text_with_links
except ImportError:
    # Fallback if import fails
    def enrich_text_with_links(text):
        return text

def extract_text_with_hyperlinks_from_docx(file_content: bytes) -> str:
    """
    Extract plain text from DOCX document and enrich with hardcoded playbook links.
    Simplified version: No complex hyperlink extraction, just text + hardcoded links.
    """
    try:
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        # Extract plain text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract plain text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                
                if row_texts:
                    text_parts.append(' | '.join(row_texts))
        
        # Combine all text
        full_text = '\n'.join(text_parts)
        
        # Enrich with hardcoded playbook links (this adds the formatted links)
        enriched_text = enrich_text_with_links(full_text)
        
        return enriched_text
        
    except Exception as e:
        # Fallback to simple text extraction
        print(f"⚠️ DOCX extraction error: {str(e)}")
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            return '\n'.join(text_parts)
        except:
            return ""

