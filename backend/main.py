from fastapi import FastAPI, Depends, UploadFile, File, Form, BackgroundTasks, Request
from fastapi.responses import FileResponse, RedirectResponse, JSONResponse
from fastapi import Request as FastAPIRequest
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import text, func, or_
from typing import List, Dict, Any
import os
from pathlib import Path
import uuid
from datetime import datetime
from dotenv import load_dotenv

from models import Base, Feedback, UploadedFile, MandatoryFile, ProjectKnowledgeBaseFile, ChatMessage, Conversation, Project, get_db, engine
from schemas import (
    LoginRequest, LoginResponse, 
    LLMChatRequest, LLMChatResponse,
    FeedbackRequest, FeedbackResponse
)
from services import (
    auth_service, 
    llm_service, gemini_service
)
from services.pdf_service import pdf_service

# Load environment variables
load_dotenv()

# Run automatic migrations first (may drop/recreate tables)
try:
    from db_migrations import run_migrations
    run_migrations()
    print("[OK] Database migrations completed")
except Exception as e:
    print(f"[WARNING] Database migrations failed: {str(e)}")
    print("[INFO] Continuing startup - some features may not work correctly")

# Create database tables (creates new tables if they don't exist)
# This runs after migrations to recreate any dropped tables
Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="PM Portal Bot API",
    description="A chatbot API with LLM integration for project management",
    version="1.0.0"
)

def _get_structured_html_system_prompt() -> str:
    """System prompt that enforces structured HTML responses for the chatbot."""
    return (
        "You are a structured project assistant chatbot.\n\n"
        "Your goal is to provide clean, organized, and readable answers based on the uploaded project file content.\n"
        "Make sure your responses follow these formatting rules:\n\n"
        "1. Use HTML structure for responses.\n"
        "2. Headings → use <h3> or <h4> tags.\n"
        "3. Lists → use <ul><li> for bullet points and <ol><li> for numbered lists.\n"
        "4. For any links (e.g., URLs found in text), make them clickable using: <a href=\"URL\" target=\"_blank\">View Document</a>\n"
        "5. Avoid showing asterisks (*) or Markdown formatting.\n"
        "6. Keep responses concise, professional, and well-indented.\n"
        "7. If a section includes multiple deliverables or items, list each one separately for readability.\n\n"
        "Example formatting:\n"
        "---------------------\n"
        "<h3>Define Phase Deliverables</h3>\n"
        "<ul>\n"
        "  <li><strong>Project Plan:</strong> Identifies all phases, activities, deliverables, and milestones.</li>\n"
        "  <li><strong>Contact List:</strong> Contains team member contact information.</li>\n"
        "  <li><strong>Project Control Procedures:</strong> Describes reporting mechanisms and change control process.</li>\n"
        "</ul>\n"
        "<p>For more details, refer to <a href=\"https://example.com\" target=\"_blank\">Project Control Document</a>.</p>\n\n"
        "Now, based on the uploaded document, answer the user’s question clearly and in this structured HTML format."
    )


# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "http://localhost:3000,http://192.168.11.101:3000").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    """Health check endpoint"""
    return {"message": "PM Portal Bot API", "status": "running"}

# Static file serving endpoint for mandatory files (by ID)
@app.get("/api/mandatory-files/{file_id}/download")
async def download_mandatory_file(file_id: int, db: Session = Depends(get_db)):
    """Download a mandatory file by ID from database"""
    from fastapi.responses import Response
    
    try:
        mandatory_file = db.query(MandatoryFile).filter(
            MandatoryFile.id == file_id,
            MandatoryFile.is_active == True
        ).first()
        
        if not mandatory_file:
            return {"success": False, "error": "File not found"}
        
        # Check if file content exists in database
        if mandatory_file.file_content:
            # Serve file from database
            file_size = len(mandatory_file.file_content) if mandatory_file.file_content else 0
            print(f"[MANDATORY-DOWNLOAD] Serving file '{mandatory_file.file_name}' (ID: {file_id}) from DATABASE - Size: {file_size} bytes")
            return Response(
                content=mandatory_file.file_content,
                media_type='application/octet-stream',
                headers={
                    "Content-Disposition": f'attachment; filename="{mandatory_file.file_name}"'
                }
            )
        elif mandatory_file.file_path:
            # Fallback: Try to serve from file system (for legacy files)
            file_path = Path(mandatory_file.file_path)
            if file_path.exists():
                print(f"[MANDATORY-DOWNLOAD] Serving file '{mandatory_file.file_name}' (ID: {file_id}) from FILE SYSTEM (legacy) - Path: {file_path}")
                return FileResponse(
                    path=str(file_path),
                    filename=mandatory_file.file_name,
                    media_type='application/octet-stream'
                )
            else:
                print(f"[MANDATORY-DOWNLOAD] File '{mandatory_file.file_name}' (ID: {file_id}) not found in database or file system")
                return {"success": False, "error": "File content not found in database or file system"}
        else:
            print(f"[MANDATORY-DOWNLOAD] File '{mandatory_file.file_name}' (ID: {file_id}) has no file_content or file_path")
            return {"success": False, "error": "File content not found"}
            
    except Exception as e:
        print(f"[MANDATORY-DOWNLOAD] Error serving file ID {file_id}: {str(e)}")
        return {"success": False, "error": f"Error serving file: {str(e)}"}

# Legacy endpoint for backward compatibility
@app.get("/mandatory/{filename:path}")
async def get_mandatory_file(filename: str):
    """Serve mandatory files from backend/mandatory directory (legacy)"""
    try:
        # Get the backend directory path
        backend_dir = Path(__file__).parent
        mandatory_dir = backend_dir / "mandatory"
        file_path = mandatory_dir / filename
        
        # Security: Ensure file is within mandatory directory (prevent path traversal)
        try:
            file_path.resolve().relative_to(mandatory_dir.resolve())
        except ValueError:
            return {"success": False, "error": "Invalid file path"}
        
        # Check if file exists
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        # Return file with appropriate media type
        return FileResponse(
            path=str(file_path),
            filename=filename,
            media_type='application/octet-stream'
        )
    except Exception as e:
        return {"success": False, "error": f"Error serving file: {str(e)}"}

# Mandatory Files Management Endpoints
@app.get("/api/mandatory-files")
async def get_mandatory_files(db: Session = Depends(get_db), include_content: bool = False):
    """Get all active mandatory files"""
    try:
        files = db.query(MandatoryFile).filter(MandatoryFile.is_active == True).order_by(MandatoryFile.uploaded_at.desc()).all()
        return {
            "success": True,
            "files": [
                {
                    "id": f.id,
                    "file_name": f.file_name,
                    "file_type": f.file_type,
                    "file_size": f.file_size,
                    "uploaded_by": f.uploaded_by,
                    "uploaded_at": f.uploaded_at.isoformat() if f.uploaded_at else None,
                    "description": f.description,
                    "extracted_text": f.extracted_text if include_content else None,  # Only include if requested
                    "has_content": bool(f.extracted_text)  # Indicate if content exists
                }
                for f in files
            ]
        }
    except Exception as e:
        return {"success": False, "error": f"Error fetching mandatory files: {str(e)}"}

@app.post("/api/mandatory-files/upload")
async def upload_mandatory_file(
    file: UploadFile = File(...),
    uploaded_by: str = Form(None),
    description: str = Form(None),
    background_tasks: BackgroundTasks = None,
    db: Session = Depends(get_db)
):
    """Upload a new mandatory file"""
    import uuid
    from datetime import datetime
    
    try:
        # Validate file type
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        if file_extension not in ['pdf', 'docx', 'txt', 'doc', 'xlsx', 'xls', 'pptx', 'ppt']:
            return {
                "success": False,
                "error": f"Invalid file type. Supported formats: PDF, DOCX, TXT, XLSX, PPTX. Got: {file_extension}"
            }
        
        # Read file content into memory (stored in database, not file system)
        file_content = await file.read()
        file_size = len(file_content)
        
        # Extract text content based on file type
        extracted_text = ""
        try:
            if file_extension == 'pdf':
                result = pdf_service.extract_text_from_pdf(file_content)
                if result['success']:
                    extracted_text = result['text']
            elif file_extension in ['docx', 'doc']:
                # Use enhanced docx extraction to preserve hyperlinks
                try:
                    from services.docx_extraction_helper import extract_text_with_hyperlinks_from_docx
                    extracted_text = extract_text_with_hyperlinks_from_docx(file_content)
                except Exception as e:
                    # Fallback to simple extraction if enhanced extraction fails
                    print(f"[MANDATORY-UPLOAD] Enhanced DOCX extraction failed, using fallback: {str(e)}")
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(file_content))
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                    extracted_text = '\n'.join(text_parts)
            elif file_extension == 'txt':
                extracted_text = file_content.decode('utf-8', errors='ignore')
            elif file_extension in ['xlsx', 'xls']:
                import io
                from openpyxl import load_workbook
                workbook = load_workbook(filename=io.BytesIO(file_content), data_only=True)
                lines = []
                for sheet in workbook.worksheets:
                    lines.append(f"Sheet: {sheet.title}")
                    for row in sheet.iter_rows(values_only=True):
                        cells = [str(cell) for cell in row if cell is not None]
                        if cells:
                            lines.append("\t".join(cells))
                extracted_text = "\n".join(lines)
            elif file_extension in ['pptx', 'ppt']:
                # PowerPoint files - basic extraction (can be enhanced later)
                extracted_text = f"[PowerPoint file: {file.filename}]"
        except Exception as e:
            print(f"Warning: Could not extract text from file: {e}")
            extracted_text = None  # Store None if extraction fails
        
        # Save to database (file content stored in DB, not file system)
        user_email = uploaded_by or "anonymous"
        mandatory_file = MandatoryFile(
            file_name=file.filename,
            file_type=file_extension,
            file_path=None,  # No longer using file system storage
            file_content=file_content,  # Store file content in database
            file_size=file_size,
            uploaded_by=user_email,
            description=description,
            is_active=True,
            extracted_text=extracted_text  # Store extracted text content
        )
        
        db.add(mandatory_file)
        db.commit()
        db.refresh(mandatory_file)
        
        print(f"[MANDATORY-UPLOAD] Saved file '{file.filename}' (ID: {mandatory_file.id}) to DATABASE - Size: {file_size} bytes, User: {user_email}")
        
        return {
            "success": True,
            "file_id": mandatory_file.id,
            "file_name": mandatory_file.file_name,
            "message": f"File '{file.filename}' uploaded successfully"
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error uploading file: {str(e)}"
        }

@app.delete("/api/mandatory-files/{file_id}")
async def delete_mandatory_file(file_id: int, db: Session = Depends(get_db)):
    """Delete a mandatory file permanently from database"""
    try:
        mandatory_file = db.query(MandatoryFile).filter(MandatoryFile.id == file_id).first()
        
        if not mandatory_file:
            return {"success": False, "error": "File not found"}
        
        file_name = mandatory_file.file_name
        
        # Remove from all users' project knowledge bases first
        db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.mandatory_file_id == file_id
        ).delete()
        
        # Hard delete - permanently remove from database (file content is in DB, no file system cleanup needed)
        db.delete(mandatory_file)
        db.commit()
        
        return {
            "success": True,
            "message": f"File '{file_name}' deleted successfully"
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error deleting file: {str(e)}"
        }

# Project Knowledge Base Endpoints
@app.post("/api/project-knowledge-base/add")
async def add_project_knowledge_base_file(
    file_id: int = Form(...),
    user_email: str = Form(...),
    db: Session = Depends(get_db)
):
    """Add a file to user's project knowledge base"""
    try:
        # Verify file exists and is active
        mandatory_file = db.query(MandatoryFile).filter(
            MandatoryFile.id == file_id,
            MandatoryFile.is_active == True
        ).first()
        
        if not mandatory_file:
            return {
                "success": False,
                "error": "File not found or inactive"
            }
        
        # Check if already exists
        existing = db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.user_email == user_email,
            ProjectKnowledgeBaseFile.mandatory_file_id == file_id
        ).first()
        
        if existing:
            # File already in knowledge base, but still ensure it's indexed in Pinecone
            print(f"[KNOWLEDGE-BASE] File {file_id} already in knowledge base, ensuring Pinecone index exists...")
            should_index = True  # Still index/reindex to ensure Pinecone is up to date
        else:
            # Add to knowledge base
            knowledge_base_file = ProjectKnowledgeBaseFile(
                user_email=user_email,
                mandatory_file_id=file_id
            )
            
            db.add(knowledge_base_file)
            db.commit()
            db.refresh(knowledge_base_file)
            should_index = True
        
        # Index file to Pinecone (separate index per file)
        # This runs whether file is new or existing to ensure Pinecone is synced
        if should_index:
            try:
                from services.pinecone_service import pinecone_service
                from services.chunking_service import chunking_service
                from services.embedding_service import embedding_service
                
                print(f"[PINECONE] Indexing mandatory file {file_id} ({mandatory_file.file_name}) to Pinecone...")
                
                # Check if file has extracted text
                if not mandatory_file.extracted_text:
                    print(f"[PINECONE] File {file_id} has no extracted_text, skipping Pinecone indexing")
                else:
                    # Create index for this file
                    index_result = pinecone_service.create_index_for_file(
                        file_id=file_id,
                        file_name=mandatory_file.file_name
                    )
                    
                    if not index_result.get("success"):
                        print(f"[PINECONE] Failed to create index: {index_result.get('error')}")
                    else:
                        # Chunk text (400 chars, 100 overlap)
                        chunks = chunking_service.chunk_text_by_characters(
                            text=mandatory_file.extracted_text,
                            chunk_size=400,
                            chunk_overlap=100,
                            metadata={
                                "file_id": file_id,
                                "file_name": mandatory_file.file_name,
                                "file_type": mandatory_file.file_type or "unknown"
                            }
                        )
                        
                        if chunks:
                            # Generate embeddings
                            chunk_texts = [chunk["text"] for chunk in chunks]
                            embeddings = embedding_service.embed(chunk_texts)
                            
                            # Index to Pinecone
                            index_chunks_result = pinecone_service.index_file_chunks(
                                file_id=file_id,
                                file_name=mandatory_file.file_name,
                                chunks=chunks,
                                embeddings=embeddings
                            )
                            
                            if index_chunks_result.get("success"):
                                print(f"[PINECONE] Successfully indexed {index_chunks_result.get('chunks_indexed', 0)} chunks for file {file_id}")
                            else:
                                print(f"[PINECONE] Failed to index chunks: {index_chunks_result.get('error')}")
                        else:
                            print(f"[PINECONE] No chunks created for file {file_id}")
            except Exception as e:
                print(f"[PINECONE] Error during Pinecone indexing: {str(e)}")
                import traceback
                print(traceback.format_exc())
                # Don't fail the request if Pinecone indexing fails
        
        return {
            "success": True,
            "message": f"File '{mandatory_file.file_name}' added to knowledge base" + (" (reindexed)" if existing else ""),
            "file_id": file_id
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error adding file to knowledge base: {str(e)}"
        }

@app.delete("/api/project-knowledge-base/remove")
async def remove_project_knowledge_base_file(
    file_id: int,
    user_email: str,
    db: Session = Depends(get_db)
):
    """Remove a file from user's project knowledge base"""
    try:
        knowledge_base_file = db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.user_email == user_email,
            ProjectKnowledgeBaseFile.mandatory_file_id == file_id
        ).first()
        
        if not knowledge_base_file:
            return {
                "success": False,
                "error": "File not found in knowledge base"
            }
        
        # Get file info before deletion
        mandatory_file = db.query(MandatoryFile).filter(
            MandatoryFile.id == file_id
        ).first()
        
        db.delete(knowledge_base_file)
        db.commit()
        
        # Delete Pinecone index for this file
        if mandatory_file:
            try:
                from services.pinecone_service import pinecone_service
                delete_result = pinecone_service.delete_index(
                    file_id=file_id,
                    file_name=mandatory_file.file_name
                )
                if delete_result.get("success"):
                    print(f"[PINECONE] Deleted index for file {file_id} ({mandatory_file.file_name})")
                else:
                    print(f"[PINECONE] Failed to delete index: {delete_result.get('error')}")
            except Exception as e:
                print(f"[PINECONE] Error deleting index: {str(e)}")
        
        return {
            "success": True,
            "message": "File removed from knowledge base",
            "file_id": file_id
        }
        
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error removing file from knowledge base: {str(e)}"
        }

@app.get("/api/project-knowledge-base")
async def get_project_knowledge_base(user_email: str, db: Session = Depends(get_db)):
    """Get all files in user's project knowledge base"""
    try:
        knowledge_base_files = db.query(ProjectKnowledgeBaseFile).filter(
            ProjectKnowledgeBaseFile.user_email == user_email
        ).all()
        
        file_ids = [kb_file.mandatory_file_id for kb_file in knowledge_base_files]
        
        # Get file details
        files = db.query(MandatoryFile).filter(
            MandatoryFile.id.in_(file_ids),
            MandatoryFile.is_active == True
        ).all()
        
        return {
            "success": True,
            "file_ids": file_ids,
            "files": [
                {
                    "id": f.id,
                    "file_name": f.file_name,
                    "file_type": f.file_type,
                    "file_size": f.file_size,
                    "uploaded_by": f.uploaded_by,
                    "uploaded_at": f.uploaded_at.isoformat() if f.uploaded_at else None
                }
                for f in files
            ]
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Error fetching knowledge base: {str(e)}"
        }

@app.post("/api/project-knowledge-base/reindex-all")
async def reindex_all_knowledge_base_files(db: Session = Depends(get_db)):
    """
    Reindex all files in the knowledge base to Pinecone.
    Useful for indexing files that were added before Pinecone integration.
    """
    try:
        from services.pinecone_service import pinecone_service
        from services.chunking_service import chunking_service
        from services.embedding_service import embedding_service
        
        # Get all knowledge base files
        knowledge_base_files = db.query(ProjectKnowledgeBaseFile).all()
        
        if not knowledge_base_files:
            return {
                "success": False,
                "error": "No files in knowledge base to reindex"
            }
        
        # Get unique file IDs
        file_ids = list(set([kb_file.mandatory_file_id for kb_file in knowledge_base_files]))
        
        # Get mandatory files
        mandatory_files = db.query(MandatoryFile).filter(
            MandatoryFile.id.in_(file_ids),
            MandatoryFile.is_active == True,
            MandatoryFile.extracted_text.isnot(None),
            MandatoryFile.extracted_text != ""
        ).all()
        
        results = []
        success_count = 0
        error_count = 0
        
        for mandatory_file in mandatory_files:
            try:
                print(f"[REINDEX] Indexing file {mandatory_file.id} ({mandatory_file.file_name})...")
                
                # Create index
                index_result = pinecone_service.create_index_for_file(
                    file_id=mandatory_file.id,
                    file_name=mandatory_file.file_name
                )
                
                if not index_result.get("success"):
                    error_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": False,
                        "error": index_result.get("error", "Failed to create index")
                    })
                    continue
                
                # Chunk text
                chunks = chunking_service.chunk_text_by_characters(
                    text=mandatory_file.extracted_text,
                    chunk_size=400,
                    chunk_overlap=100,
                    metadata={
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "file_type": mandatory_file.file_type or "unknown"
                    }
                )
                
                if not chunks:
                    error_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": False,
                        "error": "No chunks created"
                    })
                    continue
                
                # Generate embeddings
                chunk_texts = [chunk["text"] for chunk in chunks]
                embeddings = embedding_service.embed(chunk_texts)
                
                # Index to Pinecone
                index_chunks_result = pinecone_service.index_file_chunks(
                    file_id=mandatory_file.id,
                    file_name=mandatory_file.file_name,
                    chunks=chunks,
                    embeddings=embeddings
                )
                
                if index_chunks_result.get("success"):
                    success_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": True,
                        "chunks_indexed": index_chunks_result.get("chunks_indexed", 0),
                        "index_name": index_result.get("index_name")
                    })
                    print(f"[REINDEX] Successfully indexed {index_chunks_result.get('chunks_indexed', 0)} chunks for file {mandatory_file.id}")
                else:
                    error_count += 1
                    results.append({
                        "file_id": mandatory_file.id,
                        "file_name": mandatory_file.file_name,
                        "success": False,
                        "error": index_chunks_result.get("error", "Failed to index chunks")
                    })
                    
            except Exception as e:
                error_count += 1
                results.append({
                    "file_id": mandatory_file.id,
                    "file_name": mandatory_file.file_name,
                    "success": False,
                    "error": str(e)
                })
                print(f"[REINDEX] Error indexing file {mandatory_file.id}: {str(e)}")
        
        return {
            "success": True,
            "total_files": len(mandatory_files),
            "success_count": success_count,
            "error_count": error_count,
            "results": results
        }
        
    except Exception as e:
        import traceback
        print(f"[REINDEX] Error: {str(e)}")
        print(traceback.format_exc())
        return {
            "success": False,
            "error": f"Error reindexing files: {str(e)}"
        }

# Authentication endpoints
@app.get("/api/auth/google/url")
async def get_google_auth_url(prompt: str = None):
    """Get Google OAuth URL"""
    try:
        print(f"[AUTH] Generating Google OAuth URL...")
        print(f"[AUTH] GOOGLE_CLIENT_ID: {'Set' if os.getenv('GOOGLE_CLIENT_ID') else 'NOT SET'}")
        print(f"[AUTH] GOOGLE_REDIRECT_URI: {os.getenv('GOOGLE_REDIRECT_URI', 'NOT SET')}")
        auth_url = auth_service.get_google_auth_url(prompt=prompt)
        print(f"[AUTH] Generated auth URL: {auth_url[:100]}...")
        return {"auth_url": auth_url}
    except ValueError as e:
        print(f"[AUTH] ValueError: {str(e)}")
        return {"error": str(e), "auth_url": None}
    except Exception as e:
        print(f"[AUTH] Exception: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return {"error": f"Failed to generate Google OAuth URL: {str(e)}", "auth_url": None}

from pydantic import BaseModel

class GoogleCallbackRequest(BaseModel):
    code: str

@app.post("/api/auth/google/callback")
async def google_auth_callback(request: GoogleCallbackRequest, db: Session = Depends(get_db)):
    """Handle Google OAuth callback"""
    try:
        print(f"[AUTH] Received Google callback with code: {request.code[:20] if request.code else 'None'}...")
        result = auth_service.authenticate_user(request.code, db)
        if result.success:
            print(f"[AUTH] Authentication successful for user: {result.user.get('email', 'Unknown') if result.user else 'None'}")
        else:
            print(f"[AUTH] Authentication failed: {result.message}")
        return result
    except Exception as e:
        print(f"[AUTH] Callback error: {str(e)}")
        import traceback
        print(traceback.format_exc())
        from schemas import LoginResponse
        return LoginResponse(
            success=False,
            session_id="",
            user=None,
            message=f"Authentication error: {str(e)}"
        )

@app.post("/api/auth/login", response_model=LoginResponse)
async def login(request: LoginRequest):
    """Simulate Google OAuth login (for demo)"""
    return auth_service.simulate_login(request)

@app.post("/api/auth/logout")
async def logout():
    """Simulate logout"""
    return auth_service.simulate_logout()

# Session cookie helper functions
def _set_session_cookie(response, session_id: str):
    """Set session cookie on response"""
    response.set_cookie(
        key="session_id",
        value=session_id,
        max_age=86400,  # 24 hours
        httponly=True,
        secure=False,  # Set to True in production with HTTPS
        samesite="lax"
    )

def _clear_session_cookie(response):
    """Clear session cookie on response"""
    response.delete_cookie(key="session_id")

@app.get("/api/auth/session")
async def get_session(
    request: FastAPIRequest,
    db: Session = Depends(get_db)
):
    """Get current session from cookie"""
    from models import Session as UserSession
    
    try:
        session_id = request.cookies.get("session_id")
        if not session_id:
            return JSONResponse(
                status_code=200,
                content={
                    "success": False,
                    "session_id": None,
                    "user": None,
                    "message": "No session cookie found"
                }
            )
        
        # Validate session from database
        session = db.query(UserSession).filter(
            UserSession.id == session_id,
            UserSession.is_active == True
        ).first()
        
        if not session or (session.expires_at and session.expires_at < datetime.now()):
            return JSONResponse(
                status_code=200,
                content={
                    "success": False,
                    "session_id": None,
                    "user": None,
                    "message": "Session expired or invalid"
                }
            )
        
        # Get user info
        from models import User
        user = db.query(User).filter(User.id == session.user_id).first()
        
        if not user:
            return JSONResponse(
                status_code=200,
                content={
                    "success": False,
                    "session_id": None,
                    "user": None,
                    "message": "User not found"
                }
            )
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "session_id": session_id,
                "user": {
                    "id": user.id,
                    "email": user.email,
                    "name": user.name,
                    "google_id": user.google_id
                },
                "message": "Session valid"
            }
        )
    except Exception as e:
        print(f"[AUTH] Session check error: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JSONResponse(
            status_code=200,
            content={
                "success": False,
                "session_id": None,
                "user": None,
                "message": f"Session check error: {str(e)}"
            }
        )

@app.get("/api/auth/login-by-email")
async def login_by_email(
    request: FastAPIRequest,
    email: str = None, 
    name: str = None, 
    redirect: str = None,
    format: str = None,  # 'json' to force JSON response
    db: Session = Depends(get_db)
):
    """Login user directly by email address (for development/testing).
    
    This endpoint accepts email as a query parameter and automatically:
    1. Finds or creates the user by email
    2. Creates a session
    3. Sets session cookie
    4. Redirects to frontend (browser) or returns JSON (API client)
    
    Usage:
        Browser: GET /api/auth/login-by-email?email=user@forsysinc.com
        API: GET /api/auth/login-by-email?email=user@forsysinc.com&format=json
    
    Note: Only @forsysinc.com email addresses are allowed.
    
    Parameters:
        - email: User email address (required)
        - name: User name (optional)
        - redirect: If 'true', redirects to frontend home page (optional, auto-detected)
        - format: If 'json', forces JSON response (for API clients like Postman)
    
    Note: This bypasses Google OAuth. Use only for trusted environments.
    """
    if not email:
        return JSONResponse(
            status_code=400,
            content={
                "success": False,
                "session_id": None,
                "user": None,
                "message": "Email parameter is required. Usage: /api/auth/login-by-email?email=user@forsysinc.com"
            }
        )
    
    # Validate email domain - only allow @forsysinc.com emails
    email_lower = email.strip().lower()
    allowed_domain = "@forsysinc.com"
    if not email_lower.endswith(allowed_domain):
        return JSONResponse(
            status_code=403,
            content={
                "success": False,
                "session_id": None,
                "user": None,
                "message": f"Access restricted. Only {allowed_domain} email addresses are allowed."
            }
        )
    
    try:
        print(f"[AUTH] Login by email requested: {email}")
        result = auth_service.login_by_email(email=email, name=name, db=db)
        
        if result.success:
            print(f"[AUTH] Email login successful for: {email}")
            
            # Determine if we should redirect or return JSON
            # Check Accept header to detect browser vs API client
            accept_header = request.headers.get("Accept", "")
            is_browser_request = (
                "text/html" in accept_header or 
                accept_header == "" or
                "*/*" in accept_header
            )
            
            # Force JSON if format=json parameter is set (for API clients like Postman)
            should_return_json = (format and format.lower() == 'json') or (
                not is_browser_request and 
                not (redirect and redirect.lower() == 'true')
            )
            
            # If redirect is explicitly requested OR it's a browser request, redirect to frontend
            if not should_return_json:
                frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
                # Redirect to email-login route which will handle the session properly
                # This ensures localStorage is set and user state is properly initialized
                redirect_url = f"{frontend_url}/email-login?email={email}"
                if name:
                    redirect_url += f"&name={name}"
                response = RedirectResponse(url=redirect_url, status_code=302)
                _set_session_cookie(response, result.session_id)
                print(f"[AUTH] Redirecting to frontend email-login: {redirect_url}")
                return response
            else:
                # Return JSON response for API calls (like Postman)
                response = JSONResponse(content=result.dict())
                _set_session_cookie(response, result.session_id)
                return response
        else:
            print(f"[AUTH] Email login failed: {result.message}")
            response = JSONResponse(content=result.dict())
            _clear_session_cookie(response)
            return response
    except Exception as e:
        print(f"[AUTH] Email login error: {str(e)}")
        import traceback
        print(traceback.format_exc())
        error_response = LoginResponse(
            success=False,
            session_id="",
            user=None,
            message=f"Email login error: {str(e)}"
        )
        response = JSONResponse(content=error_response.dict())
        _clear_session_cookie(response)
        return response

# SOW upload endpoint: supports both DOCX and PDF files; prints content to console and returns raw text
@app.post("/api/upload/sow")
async def upload_sow_document(file: UploadFile = File(...), user_email: str = Form(None)):
    """Upload SOW document (DOCX or PDF), print raw text to server console, return raw text back."""
    try:
        file_extension = file.filename.lower().split('.')[-1]
        
        if file_extension not in ['docx', 'pdf']:
            return {"success": False, "error": "Invalid file type. Please upload a .docx or .pdf file."}

        file_content = await file.read()
        raw_text = ""

        if file_extension == 'docx':
            # Convert DOCX to HTML preserving formatting
            from docx import Document
            import io
            import re

            doc = Document(io.BytesIO(file_content))
            html_parts = []
            
            # Process paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    # Check if paragraph has heading style
                    if p.style.name.startswith('Heading'):
                        level = p.style.name.replace('Heading ', '')
                        html_parts.append(f'<h{level}>{p.text.strip()}</h{level}>')
                    else:
                        # Check for bold/italic runs
                        para_html = ""
                        for run in p.runs:
                            text = run.text
                            if run.bold:
                                text = f"<strong>{text}</strong>"
                            if run.italic:
                                text = f"<em>{text}</em>"
                            para_html += text
                        
                        if para_html.strip():
                            html_parts.append(f'<p>{para_html}</p>')
            
            # Process tables
            for table in doc.tables:
                html_parts.append('<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">')
                for row_idx, row in enumerate(table.rows):
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Check if it's a header row (usually first row)
                            tag = 'th' if row_idx == 0 else 'td'
                            html_parts.append(f'<{tag} style="padding: 8px; border: 1px solid #ddd;">{cell_text}</{tag}>')
                        else:
                            tag = 'th' if row_idx == 0 else 'td'
                            html_parts.append(f'<{tag} style="padding: 8px; border: 1px solid #ddd;">&nbsp;</{tag}>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')
            
            # Join HTML parts
            html_content = '\n'.join(html_parts)
            
            # Also create raw text version for backward compatibility
            raw_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            
        elif file_extension == 'pdf':
            # Extract raw text from PDF using our PDF service
            pdf_result = pdf_service.extract_text_from_pdf(file_content)
            
            if not pdf_result['success']:
                return {"success": False, "error": f"Error extracting text from PDF: {pdf_result['error']}"}
            
            raw_text = pdf_result['text'].strip()
            print(f"[PDF EXTRACTION] Used method: {pdf_result.get('method', 'Unknown')}")

        if not raw_text:
            return {"success": False, "error": "No text content found in the document."}

        # Print to terminal/console
        print("=" * 80)
        print(f"[SOW UPLOAD] File: {file.filename} ({file_extension.upper()})")
        print("— Raw Document Text —")
        print(raw_text if raw_text else "<No text content>")
        print("=" * 80)

        # Return both HTML and raw text
        response_data = {
            "rawText": raw_text, 
            "fileName": file.filename, 
            "fileType": file_extension
        }
        
        # Add HTML content for DOCX files
        if file_extension == 'docx' and 'html_content' in locals():
            response_data["htmlContent"] = html_content
        
        return {"success": True, "data": response_data}
    except Exception as e:
        return {"success": False, "error": f"Error processing SOW file: {str(e)}"}

# LLM and Gemini endpoints
@app.post("/api/llm/chat", response_model=LLMChatResponse)
async def llm_chat(request: LLMChatRequest):
    """Mock LLM chat endpoint"""
    return llm_service.chat(request)



@app.post("/api/gemini/chat")
async def gemini_chat(request: dict):
    """Gemini chat endpoint"""
    from services.gemini_service import gemini_service
    messages = request.get("messages", []) or []

    # Prepend structured HTML system prompt if not already included
    system_prompt = _get_structured_html_system_prompt()
    if not (messages and messages[0].get("role") == "system"):
        messages = [{"role": "system", "content": system_prompt}] + messages
    else:
        # If a system message exists, combine ours ahead of it
        messages = [{"role": "system", "content": system_prompt}] + messages

    return gemini_service.chat(messages, request.get("max_tokens", 3000))

# Background task for indexing files in Pinecone
def index_file_background(file_id: int, text: str, source_filename: str, file_type: str, uploaded_by: str, uploaded_at):
    """
    Background task to index a file in Pinecone.
    This runs asynchronously and doesn't block the upload response.
    """
    try:
        from services.pinecone_service import pinecone_service
        from services.chunking_service import chunking_service
        from services.embedding_service import embedding_service
        from models import SessionLocal
        from datetime import datetime
        import traceback
        
        # Create a new database session for the background task
        db = SessionLocal()
        try:
            # Get the uploaded file record
            from models import UploadedFile
            uploaded_file = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
            
            if not uploaded_file:
                print(f"[BACKGROUND INDEX] File {file_id} not found in database")
                return
            
            # Use extracted_text from database if passed text is empty or None
            text_to_index = text
            if not text_to_index or not text_to_index.strip():
                text_to_index = uploaded_file.extracted_text or ""
                print(f"[BACKGROUND INDEX] File {file_id}: Using extracted_text from database (length: {len(text_to_index) if text_to_index else 0})")
            
            # Check if we have text to index
            if not text_to_index or not text_to_index.strip():
                print(f"[BACKGROUND INDEX] File {file_id}: No text content to index (empty extracted_text)")
                uploaded_file.indexing_status = "error"
                db.commit()
                return
            
            # Handle datetime conversion if needed
            upload_datetime = uploaded_at
            if isinstance(upload_datetime, str):
                try:
                    # Try parsing ISO format datetime string
                    upload_datetime = datetime.fromisoformat(upload_datetime.replace('Z', '+00:00'))
                except:
                    upload_datetime = uploaded_file.upload_time or datetime.utcnow()
            elif upload_datetime is None:
                upload_datetime = uploaded_file.upload_time or datetime.utcnow()
            
            # Use values from database if not provided
            source_filename = source_filename or uploaded_file.file_name or "unknown"
            file_type = file_type or uploaded_file.file_type or "unknown"
            uploaded_by = uploaded_by or uploaded_file.uploaded_by or "anonymous"
            
            print(f"[BACKGROUND INDEX] Starting indexing for file {file_id} ({source_filename}), text length: {len(text_to_index)}")
            
            # Create Pinecone index for this file
            pinecone_index_result = pinecone_service.create_index_for_file(
                file_id=file_id,
                file_name=source_filename
            )
            
            if not pinecone_index_result.get("success"):
                error_msg = pinecone_index_result.get("error", "Unknown error creating Pinecone index")
                uploaded_file.indexing_status = "error"
                db.commit()
                print(f"[BACKGROUND INDEX] File {file_id} indexing failed: {error_msg}")
                return
            
            # Chunk text for Pinecone (400 chars, 100 overlap)
            chunks = chunking_service.chunk_text_by_characters(
                text=text_to_index,
                chunk_size=400,
                chunk_overlap=100,
                metadata={
                    "file_id": file_id,
                    "file_name": source_filename,
                    "file_type": file_type,
                    "uploaded_by": uploaded_by,
                    "uploaded_at": upload_datetime.isoformat() if upload_datetime else None
                }
            )
            
            if not chunks:
                uploaded_file.indexing_status = "error"
                db.commit()
                print(f"[BACKGROUND INDEX] File {file_id}: No chunks generated for Pinecone indexing")
                return
            
            # Generate embeddings for chunks
            chunk_texts = [chunk["text"] for chunk in chunks]
            embeddings = embedding_service.embed(chunk_texts)
            
            # Index chunks in Pinecone
            index_chunks_result = pinecone_service.index_file_chunks(
                file_id=file_id,
                file_name=source_filename,
                chunks=chunks,
                embeddings=embeddings
            )
            
            if index_chunks_result.get("success"):
                uploaded_file.indexing_status = "indexed"
                db.commit()
                print(f"[BACKGROUND INDEX] File {file_id} indexed successfully in Pinecone: {index_chunks_result.get('chunks_indexed', 0)} chunks")
            else:
                error_msg = index_chunks_result.get('error', 'Unknown error during Pinecone indexing')
                uploaded_file.indexing_status = "error"
                db.commit()
                print(f"[BACKGROUND INDEX] File {file_id} indexing failed: {error_msg}")
                print(f"[BACKGROUND INDEX] File {file_id} error details: {traceback.format_exc()}")
            
        except Exception as e:
            # Update status to error if indexing fails
            error_msg = str(e)
            error_trace = traceback.format_exc()
            print(f"[BACKGROUND INDEX] File {file_id} indexing exception: {error_msg}")
            print(f"[BACKGROUND INDEX] File {file_id} exception traceback:\n{error_trace}")
            
            try:
                uploaded_file = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
                if uploaded_file:
                    uploaded_file.indexing_status = "error"
                    db.commit()
            except Exception as db_error:
                print(f"[BACKGROUND INDEX] Failed to update error status: {str(db_error)}")
        finally:
            db.close()
    except Exception as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()
        print(f"[BACKGROUND INDEX] Critical error indexing file {file_id}: {error_msg}")
        print(f"[BACKGROUND INDEX] Critical error traceback:\n{error_trace}")

# Helper function to process a single file
async def process_single_file(
    file: UploadFile,
    uploaded_by: str,
    db: Session,
    background_tasks: BackgroundTasks
) -> dict:
    """
    Process a single file: validate, save, extract text, save to DB, and queue indexing.
    Returns a dict with success status and file information or error.
    """
    import uuid
    from pathlib import Path
    
    try:
        # Validate file type
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        if file_extension not in ['pdf', 'docx', 'txt', 'doc', 'xlsx']:
            return {
                "success": False,
                "error": f"Invalid file type. Supported formats: PDF, DOCX, TXT, XLSX. Got: {file_extension}",
                "file_name": file.filename
            }
        
        # Create uploads directory if it doesn't exist
        upload_dir = Path("uploads")
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        unique_filename = f"{uuid.uuid4()}_{file.filename}"
        file_path = upload_dir / unique_filename
        
        # Save file to disk
        file_content = await file.read()
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        # Extract text based on file type
        extracted_text = ""
        if file_extension == 'pdf':
            result = pdf_service.extract_text_from_pdf(file_content)
            if result['success']:
                extracted_text = result['text']
            else:
                return {
                    "success": False,
                    "error": f"Failed to extract text from PDF: {result.get('error', 'Unknown error')}",
                    "file_name": file.filename
                }
        elif file_extension in ['docx', 'doc']:
            # Use enhanced docx extraction to preserve hyperlinks
            try:
                from services.docx_extraction_helper import extract_text_with_hyperlinks_from_docx
                extracted_text = extract_text_with_hyperlinks_from_docx(file_content)
            except Exception as e:
                # Fallback to simple extraction if enhanced extraction fails
                print(f"[UPLOAD] Enhanced DOCX extraction failed for {file.filename}, using fallback: {str(e)}")
                from docx import Document
                import io
                doc = Document(io.BytesIO(file_content))
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                extracted_text = '\n'.join(text_parts)
        elif file_extension == 'txt':
            extracted_text = file_content.decode('utf-8', errors='ignore')
        elif file_extension == 'xlsx':
            # Extract text from Excel using openpyxl (cells by rows)
            try:
                import io
                from openpyxl import load_workbook
                workbook = load_workbook(filename=io.BytesIO(file_content), data_only=True)
                lines = []
                
                for sheet in workbook.worksheets:
                    sheet_lines = [f"Sheet: {sheet.title}"]
                    row_count = 0
                    for row in sheet.iter_rows(values_only=True):
                        # Filter out None, empty strings, and convert to strings
                        cells = []
                        for cell in row:
                            if cell is not None:
                                cell_str = str(cell).strip()
                                if cell_str:  # Only add non-empty cells
                                    cells.append(cell_str)
                        
                        if cells:  # Only add rows with content
                            sheet_lines.append("\t".join(cells))
                            row_count += 1
                    
                    # Only add sheet if it has content
                    if row_count > 0:
                        lines.extend(sheet_lines)
                
                extracted_text = "\n".join(lines) if lines else ""
                
                # Ensure we have some text content
                if not extracted_text or not extracted_text.strip():
                    extracted_text = f"[XLSX file: {file.filename}]\nThis file contains spreadsheet data but no extractable text content was found. The file may contain only formulas, images, or empty cells."
                    print(f"[UPLOAD] XLSX file {file.filename} has no extractable text, using placeholder")
                else:
                    print(f"[UPLOAD] XLSX file {file.filename} extracted {len(extracted_text)} characters")
                    
            except Exception as e:
                print(f"[UPLOAD] XLSX extraction error for {file.filename}: {str(e)}")
                import traceback
                print(f"[UPLOAD] XLSX extraction traceback:\n{traceback.format_exc()}")
                # Don't fail the upload, use placeholder text
                extracted_text = f"[XLSX file: {file.filename}]\nError extracting text: {str(e)}\nThis file may contain only formulas, images, or be in an unsupported format."
        
        # Determine uploaded_by (use provided value or default)
        user_email = uploaded_by or "anonymous"
        
        # Save to database
        uploaded_file = UploadedFile(
            file_name=file.filename,
            file_type=file_extension,
            file_path=str(file_path),
            uploaded_by=user_email,
            status="Processed",
            extracted_text=extracted_text,
            indexing_status="pending_index"
        )
        
        db.add(uploaded_file)
        db.commit()
        db.refresh(uploaded_file)
        
        # Schedule indexing in Pinecone (always background to prevent blocking)
        if background_tasks is None:
            background_tasks = BackgroundTasks()
        
        background_tasks.add_task(
            index_file_background,
            file_id=uploaded_file.id,
            text=extracted_text,
            source_filename=uploaded_file.file_name,
            file_type=uploaded_file.file_type,
            uploaded_by=uploaded_file.uploaded_by,
            uploaded_at=uploaded_file.upload_time
        )
        print(f"[UPLOAD] File {uploaded_file.id} ({uploaded_file.file_name}) queued for Pinecone indexing")
        
        return {
            "success": True,
            "file_id": uploaded_file.id,
            "file_name": uploaded_file.file_name,
            "file_type": uploaded_file.file_type,
            "indexing_status": uploaded_file.indexing_status,
            "extracted_length": len(extracted_text),
            "message": f"File uploaded and processed successfully. Extracted {len(extracted_text)} characters. Indexing in progress..."
        }
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"[UPLOAD] Error processing file {file.filename}: {str(e)}")
        print(f"[UPLOAD] Traceback:\n{error_trace}")
        return {
            "success": False,
            "error": f"Error processing file: {str(e)}",
            "file_name": file.filename
        }

# Chatbot file upload and question endpoints
@app.post("/api/upload-file")
async def upload_chatbot_file(
    request: Request,
    uploaded_by: str = Form(None),
    background_tasks: BackgroundTasks = None,
    db: Session = Depends(get_db)
):
    """
    Upload one or multiple files for chatbot analysis
    Supports PDF, DOCX, TXT, and XLSX files
    Accepts up to 10 files in a single request
    Saves files to /uploads/ and extracts text content
    Each file is processed independently and indexed separately
    
    Backward compatible: accepts both 'file' (single) and 'files' (multiple) parameters
    """
    import uuid
    from datetime import datetime
    from pathlib import Path
    
    try:
        # Parse form data manually to handle both 'file' and 'files' parameters
        form = await request.form()
        normalized_files = []
        
        # Get uploaded_by from form if not provided as parameter
        if uploaded_by is None and 'uploaded_by' in form:
            uploaded_by = form['uploaded_by']
        
        # Check for 'files' (multiple) parameter first (new format)
        if 'files' in form:
            files_list = form.getlist('files')
            if files_list and len(files_list) > 0:
                # Filter out None values
                files_list = [f for f in files_list if f is not None]
                if files_list:
                    normalized_files = files_list
            # If 'files' exists but is empty, check 'file' parameter
            if len(normalized_files) == 0 and 'file' in form:
                file_obj = form['file']
                if file_obj:
                    normalized_files = [file_obj]
        # Fallback: check 'file' (single) parameter for backward compatibility
        elif 'file' in form:
            file_obj = form['file']
            if file_obj:
                normalized_files = [file_obj]
        
        if len(normalized_files) == 0:
            return {
                "success": False,
                "error": "No files provided. Please upload at least one file."
            }
        
        files = normalized_files
        
        # Validate number of files (max 10)
        if len(files) > 10:
            return {
                "success": False,
                "error": f"Too many files. Maximum 10 files allowed. Got {len(files)} files."
            }
        
        # Process each file independently
        results = []
        successful_files = []
        failed_files = []
        
        print(f"[UPLOAD] Processing {len(files)} file(s)...")
        
        for idx, file in enumerate(files, 1):
            print(f"[UPLOAD] Processing file {idx}/{len(files)}: {file.filename}")
            result = await process_single_file(file, uploaded_by, db, background_tasks)
            results.append(result)
            
            if result["success"]:
                successful_files.append(result)
            else:
                failed_files.append(result)
        
        # Prepare response
        total_files = len(files)
        success_count = len(successful_files)
        failed_count = len(failed_files)
        
        response = {
            "success": True,
            "total_files": total_files,
            "successful_uploads": success_count,
            "failed_uploads": failed_count,
            "files": results,
            "message": f"Processed {total_files} file(s): {success_count} successful, {failed_count} failed"
        }
        
        # If all files failed, mark overall success as False
        if success_count == 0:
            response["success"] = False
            response["error"] = "All files failed to upload"
        # If some files failed, still return success but with warning
        elif failed_count > 0:
            response["warning"] = f"{failed_count} file(s) failed to upload. Check individual file results for details."
        
        print(f"[UPLOAD] Batch upload complete: {success_count}/{total_files} successful")
        return response
        
    except Exception as e:
        db.rollback()
        import traceback
        error_trace = traceback.format_exc()
        print(f"[UPLOAD] Batch upload error: {str(e)}")
        print(f"[UPLOAD] Traceback:\n{error_trace}")
        return {
            "success": False,
            "error": f"Error processing batch upload: {str(e)}"
        }

def _router_answerer_system_prompt() -> str:
    """Returns the ROUTER + ANSWERER system prompt"""
    return """You are a reliable two-step assistant (ROUTER + ANSWERER) that receives:

  - a user question,
  - a list of candidates from the vector DB (top-K chunk hits across all files) summarized as `file_scores`, and
  - the top context chunks (documents) retrieved for each file candidate.

Your job:

  1) ROUTE: Decide which file(s) are most appropriate to answer the question using the `file_scores`.

  2) RETRIEVE/ANSWER: Use only the provided `context_chunks` (the chunks attached from Pinecone) from the selected file(s) and produce a concise, factual answer. Do NOT use external knowledge beyond light, safe inference. Always cite sources inline and in the `sources` array.

  3) OUTPUT: Return a single valid JSON object (no extra text) matching the specified schema exactly.

Routing rules (apply before answering):

  - `file_scores` gives, per file: file_name, top_chunk_score (similarity 0..1), and 1-2 line summary.
  - CONFIDENCE thresholds:
      HIGH_CONFIDENCE if top score >= 0.72
      MEDIUM_CONFIDENCE if 0.55 <= top score < 0.72
      LOW_CONFIDENCE if top score < 0.55
  - If top two files have scores within 0.03 of each other, include both files as `selected_files` (order by score).
  - If LOW_CONFIDENCE: do not hallucinate. Instead, set `"status": "LOW_CONFIDENCE"` and attempt a conservative best-effort answer using top available chunks from all candidate files, and include a `clarifying_question` suggestion for the user.

Answering rules:

  - Use ONLY the `context_chunks` provided. Do not invent facts. You may make short, reasonable inferences, but label them as "inference" if not explicitly present in chunks.
  - Keep answers concise by default (<= 200 words). If the user explicitly asks for more depth, extend.
  - Cite each factual statement that comes from a chunk using inline bracketed citations like: [Project Management Playbook.docx — chunk_12].
  - After the answer include a `Sources` section listing each chunk used with `file_name`, `chunk_id`, and `score`.
  - If no chunk contains an exact answer, explicitly say: "I couldn't find an exact answer in the provided documents; here's the best guidance based on them." Then answer and cite.

Output JSON schema (MUST match exactly):

{
  "status": "OK" | "MEDIUM_CONFIDENCE" | "LOW_CONFIDENCE",
  "selected_files": ["file1.docx", "file2.xlsx"],
  "routing_detail": {
    "top_file": "file1.docx",
    "top_score": 0.82,
    "file_scores": [ { "file_name": "...", "top_chunk_score": 0.82, "summary":"..." }, ... ]
  },
  "answer": "<the assistant's concise answer (string)>",
  "sources": [
    { "file_name": "Project Management Playbook.docx", "chunk_id": "chunk_72", "score": 0.83, "excerpt": "first 200 chars of the chunk..." },
    ...
  ],
  "confidence_explanation": "<one-line explanation of confidence and any limitations>",
  "clarifying_question": "<optional string - present only when status is LOW_CONFIDENCE or ambiguous>",
  "raw_used_chunks": [ { "file_name":"...", "chunk_id":"...", "text":"full chunk text (or truncated to X chars)", "score": 0.83 } ]
}

Formatting constraints:
  - Return only the JSON object (no surrounding commentary or markdown).
  - Truncate long chunk excerpts in `sources` to 200 chars. `raw_used_chunks` may be truncated to 2000 chars each if necessary to keep output practical."""


def _resolve_user_email(db: Session, chat_id: str, provided_email: str = None) -> str:
    """Resolve which user_email to store for a chat."""
    if provided_email:
        return provided_email
    
    try:
        existing = (
            db.query(ChatMessage.user_email)
            .filter(ChatMessage.chat_id == chat_id, ChatMessage.user_email.isnot(None))
            .order_by(ChatMessage.created_at.asc())
            .first()
        )
        if existing and existing[0]:
            return existing[0]
    except Exception:
        pass
    return None


def _save_chat_message(
    db: Session,
    chat_id: str,
    role: str,
    message: str,
    user_email: str = None
) -> None:
    """Persist chat messages by updating conversation JSON instead of creating individual records."""
    if not chat_id or not role or message is None:
        return
    
    cleaned_message = message.strip() if isinstance(message, str) else ""
    if not cleaned_message:
        return
    
    try:
        import json
        resolved_email = _resolve_user_email(db, chat_id, user_email)
        truncated_message = cleaned_message[:8000]  # Prevent excessively long entries
        
        # Find or create conversation record
        conversation = db.query(Conversation).filter(
            Conversation.chat_id == chat_id
        ).first()
        
        if conversation:
            # Update existing conversation - preserve project_id if it exists
            conversation_data = conversation.conversation_json
            if not isinstance(conversation_data, dict):
                conversation_data = json.loads(conversation_data) if isinstance(conversation_data, str) else {}
            
            messages = conversation_data.get("messages", [])
            
            # Find the last message entry or create new one
            if messages and role == "assistant" and len(messages) > 0:
                # If assistant message, update the last message entry
                last_msg = messages[-1]
                if "assistant" not in last_msg or not last_msg.get("assistant"):
                    last_msg["assistant"] = truncated_message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": "",
                        "assistant": truncated_message
                    })
            elif messages and role == "user":
                # If user message, check if last message has user or create new
                last_msg = messages[-1] if messages else None
                if last_msg and not last_msg.get("user"):
                    last_msg["user"] = truncated_message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": truncated_message,
                        "assistant": ""
                    })
            else:
                # First message
                messages.append({
                    "message_id": 1,
                    "user": truncated_message if role == "user" else "",
                    "assistant": truncated_message if role == "assistant" else ""
                })
            
            conversation_data["messages"] = messages
            conversation.conversation_json = conversation_data
            if resolved_email:
                conversation.user_email = resolved_email
            # Note: project_id is already set on the conversation, we don't need to update it
            db.commit()
        else:
            # Create new conversation
            conversation_id = str(uuid.uuid4())
            conversation_data = {
                "conversation_id": 1,  # Will be updated to actual id after save
                "messages": [{
                    "message_id": 1,
                    "user": truncated_message if role == "user" else "",
                    "assistant": truncated_message if role == "assistant" else ""
                }]
            }
            new_conversation = Conversation(
                conversation_id=conversation_id,
                chat_id=chat_id,
                user_email=resolved_email,
                conversation_json=conversation_data
            )
            db.add(new_conversation)
            db.commit()
            db.refresh(new_conversation)
            # Update conversation_id in JSON to actual database id
            conversation_data["conversation_id"] = new_conversation.id
            new_conversation.conversation_json = conversation_data
            db.commit()
    except Exception as e:
        db.rollback()
        print(f"[CHAT] Failed to save chat message: {str(e)}")


def _search_across_all_files_and_route(
    question: str,
    top_k: int = 10,
    db: Session = None
) -> Dict[str, Any]:
    """
    Search across all indexed files, group by file, calculate file scores,
    and prepare data for ROUTER + ANSWERER system.
    
    Returns:
        Dict with file_scores and context_chunks ready for LLM routing
    """
    from services.pinecone_service import pinecone_service
    from services.embedding_service import embedding_service
    from models import UploadedFile, MandatoryFile, ProjectKnowledgeBaseFile
    
    print(f"[ROUTER] Searching across all Pinecone indexes with top_k={top_k}")
    
    if db is None:
        print(f"[ROUTER] Database session not provided")
        return {"file_scores": [], "context_chunks": []}
    
    # Gather indexed uploaded files
    uploaded_files = db.query(UploadedFile).filter(UploadedFile.indexing_status == "indexed").all()
    
    # Also include knowledge base mandatory files that are indexed
    kb_file_ids = [
        kb_file.mandatory_file_id
        for kb_file in db.query(ProjectKnowledgeBaseFile).all()
    ]
    mandatory_files = []
    if kb_file_ids:
        mandatory_files = db.query(MandatoryFile).filter(MandatoryFile.id.in_(kb_file_ids)).all()
    
    if not uploaded_files and not mandatory_files:
        print(f"[ROUTER] No indexed files available for routing")
        return {"file_scores": [], "context_chunks": []}
    
    existing_indexes = pinecone_service.list_indexes()
    index_names = []
    file_info_map = {}
    
    # Add uploaded file indexes
    for file in uploaded_files:
        index_name = pinecone_service.get_index_name_for_file(file.id, file.file_name)
        if index_name in existing_indexes:
            index_names.append(index_name)
            file_info_map[index_name] = {
                "file_id": file.id,
                "file_name": file.file_name
            }
    
    # Add mandatory file indexes
    for file in mandatory_files:
        index_name = pinecone_service.get_index_name_for_file(file.id, file.file_name)
        if index_name in existing_indexes:
            index_names.append(index_name)
            file_info_map[index_name] = {
                "file_id": file.id,
                "file_name": file.file_name
            }
    
    # Deduplicate index names
    index_names = list(dict.fromkeys(index_names))
    
    if not index_names:
        print(f"[ROUTER] No Pinecone indexes available for routing")
        return {"file_scores": [], "context_chunks": []}
    
    print(f"[ROUTER] Searching across indexes: {[file_info_map[idx]['file_name'] for idx in index_names]}")
    
    query_embedding = embedding_service.embed_query(question)
    search_result = pinecone_service.search_across_indexes(
        query_embedding=query_embedding,
        index_names=index_names,
        top_k=max(3, top_k)  # ensure at least 3 per index
    )
    
    if not search_result.get("success"):
        print(f"[ROUTER] Pinecone search failed: {search_result.get('error')}")
        return {"file_scores": [], "context_chunks": []}
    
    results = search_result.get("results", [])
    if not results:
        print(f"[ROUTER] No Pinecone results found")
        return {"file_scores": [], "context_chunks": []}
    
    print(f"[ROUTER] Found {len(results)} total Pinecone results")
    
    files_dict = {}
    for result in results:
        index_name = result.get("index_name")
        metadata = result.get("metadata", {})
        file_info = file_info_map.get(index_name, {})
        file_name = file_info.get("file_name", metadata.get("file_name", "unknown"))
        file_id = file_info.get("file_id", metadata.get("file_id"))
        score = result.get("score", 0.0)
        text = metadata.get("text", "")
        
        if file_name not in files_dict:
            files_dict[file_name] = {
                "file_id": file_id,
                "chunks": [],
                "top_score": 0.0
            }
        
        files_dict[file_name]["chunks"].append({
            "score": score,
            "text": text,
            "metadata": metadata,
            "chunk_id": result.get("id", f"chunk_{metadata.get('chunk_index', '?')}"),
            "index_name": index_name
        })
        
        if score > files_dict[file_name]["top_score"]:
            files_dict[file_name]["top_score"] = score
    
    file_scores = []
    context_chunks = []
    
    for file_name, file_data in sorted(files_dict.items(), key=lambda x: x[1]["top_score"], reverse=True):
        chunks = sorted(file_data["chunks"], key=lambda x: x["score"], reverse=True)
        top_chunk = chunks[0] if chunks else None
        summary = ""
        if top_chunk:
            chunk_text = top_chunk.get("text", "")
            summary = chunk_text[:200].replace("\n", " ").strip()
            if len(chunk_text) > 200:
                summary += "..."
        
        file_scores.append({
            "file_name": file_name,
            "top_chunk_score": file_data["top_score"],
            "summary": summary
        })
        
        for chunk in chunks[:3]:
            context_chunks.append({
                "file_name": file_name,
                "chunk_id": chunk.get("chunk_id"),
                "text": chunk.get("text", ""),
                "score": chunk.get("score", 0.0)
            })
    
    print(f"[ROUTER] Prepared {len(file_scores)} file scores and {len(context_chunks)} context chunks")
    
    return {
        "file_scores": file_scores,
        "context_chunks": context_chunks
    }


@app.post("/api/ask-question")
async def ask_chatbot_question(
    question: str = Form(...),
    file_id: int = Form(None),
    file_context: str = Form(None),
    mandatory_file_ids: str = Form(None),
    chat_id: str = Form(None),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """
    Ask a question to the chatbot using vector search.
    - If file_context is provided, uses it directly (for mandatory files - backward compatibility)
    - If file_id is provided, uses Pinecone vector search to retrieve relevant chunks
    - If neither is provided, searches across all files using Pinecone indexes and ROUTER + ANSWERER system
    """
    question = (question or "").strip()
    chat_id = chat_id or str(uuid.uuid4())
    
    print(f"\n{'='*80}")
    print(f"[ASK-QUESTION] New question received")
    print(f"[ASK-QUESTION] Question: {question[:200]}...")
    print(f"[ASK-QUESTION] Parameters - file_id: {file_id}, file_context: {'Yes' if file_context else 'No'}, mandatory_file_ids: {mandatory_file_ids}")
    print(f"{'='*80}\n")
    
    try:
        from services.gemini_service import gemini_service
        from services.embedding_service import embedding_service
        from services.pinecone_service import pinecone_service
        import json
        
        context_text = ""
        use_router_answerer = False
        use_pinecone_search = False  # Initialize Pinecone search flag
        
        _save_chat_message(db, chat_id, "user", question, user_email)
        
        if file_context:
            # Use provided file_context directly (e.g., from multiple mandatory files)
            # This maintains backward compatibility for mandatory files
            context_text = file_context
            print(f"[ASK-QUESTION] Using provided file_context (PLAYBOOK/MANDATORY FILES)")
            print(f"[ASK-QUESTION] Context length: {len(context_text)} characters")
            if mandatory_file_ids:
                try:
                    import json
                    ids = json.loads(mandatory_file_ids)
                    print(f"[ASK-QUESTION] Mandatory file IDs used: {ids}")
                    # Fetch file names for these IDs
                    if ids:
                        mandatory_files = db.query(MandatoryFile).filter(MandatoryFile.id.in_(ids)).all()
                        file_names = [f.file_name for f in mandatory_files]
                        print(f"[ASK-QUESTION] Documents being used: {file_names}")
                except:
                    print(f"[ASK-QUESTION] Mandatory file IDs (raw): {mandatory_file_ids}")
            print(f"[ASK-QUESTION] Question: {question[:200]}...")
        elif file_id:
            # Use Pinecone vector search to retrieve relevant chunks
            uploaded_file = db.query(UploadedFile).filter(UploadedFile.id == file_id).first()
            if not uploaded_file:
                return {
                    "success": False,
                    "error": f"File with ID {file_id} not found.",
                    "chat_id": chat_id
                }
            
            print(f"[ASK-QUESTION] Document being used: {uploaded_file.file_name} (ID: {file_id})")
            
            # Check if file is indexed
            if uploaded_file.indexing_status != "indexed":
                # Fallback to full text if not indexed yet
                if uploaded_file.extracted_text:
                    context_text = uploaded_file.extracted_text
                    print(f"[ASK-QUESTION] File {file_id} not indexed, using full text (length: {len(context_text)} characters)")
                else:
                    return {
                        "success": False,
                        "error": f"File {file_id} is not indexed and has no extracted text available.",
                        "chat_id": chat_id
                    }
            else:
                # Use Pinecone vector search
                print(f"[ASK-QUESTION] Using Pinecone search for file_id {file_id}")
                
                index_name = pinecone_service.get_index_name_for_file(file_id, uploaded_file.file_name)
                if not pinecone_service.index_exists(index_name):
                    print(f"[ASK-QUESTION] Pinecone index {index_name} not found. Falling back to full text.")
                    if uploaded_file.extracted_text:
                        context_text = uploaded_file.extracted_text
                    else:
                        context_text = ""
                else:
                    query_embedding = embedding_service.embed_query(question)
                    search_result = pinecone_service.search_across_indexes(
                        query_embedding=query_embedding,
                        index_names=[index_name],
                        top_k=5
                    )
                    
                    results = search_result.get("results") if search_result.get("success") else []
                    
                    if results:
                        chunk_texts = []
                        for result in results:
                            metadata = result.get("metadata", {})
                            chunk_text = metadata.get("text", "")
                            chunk_index = metadata.get("chunk_index", "?")
                            file_name = metadata.get("file_name", uploaded_file.file_name)
                            score = result.get("score", 0.0)
                            chunk_texts.append(
                                f"[Chunk {chunk_index} from {file_name} (score: {score:.3f})]\n{chunk_text}"
                            )
                        
                        context_text = "\n\n---\n\n".join(chunk_texts)
                        print(f"[ASK-QUESTION] Retrieved {len(results)} relevant Pinecone chunks (total length: {len(context_text)} characters)")
                    else:
                        if uploaded_file.extracted_text:
                            context_text = uploaded_file.extracted_text
                            print(f"[ASK-QUESTION] No Pinecone results, using full text (length: {len(context_text)} characters)")
                        else:
                            context_text = ""
                            print(f"[ASK-QUESTION] No Pinecone results and no extracted text available")
        else:
            # No file_id or file_context provided - search Pinecone indexes for knowledge base files
            use_pinecone_search = True
            print(f"[ASK-QUESTION] No specific file provided, searching Pinecone knowledge base indexes...")
        
        # Build prompt with file context if available
        if use_pinecone_search:
            # Search across all Pinecone indexes for knowledge base files
            kb_context_found = False
            try:
                knowledge_base_files = db.query(ProjectKnowledgeBaseFile).all()
                
                if knowledge_base_files:
                    kb_file_ids = [kb_file.mandatory_file_id for kb_file in knowledge_base_files]
                    all_mandatory_files = db.query(MandatoryFile).filter(
                        MandatoryFile.id.in_(kb_file_ids),
                        MandatoryFile.is_active == True,
                        MandatoryFile.extracted_text.isnot(None),
                        MandatoryFile.extracted_text != ""
                    ).all()
                    
                    if all_mandatory_files:
                        index_names = []
                        file_info_map = {}
                        existing_indexes = pinecone_service.list_indexes()
                        
                        for file in all_mandatory_files:
                            index_name = pinecone_service.get_index_name_for_file(file.id, file.file_name)
                            if index_name in existing_indexes:
                                index_names.append(index_name)
                                file_info_map[index_name] = {
                                    "file_id": file.id,
                                    "file_name": file.file_name
                                }
                        
                        if index_names:
                            print(f"[PINECONE] Searching across {len(index_names)} indexes: {[file_info_map[idx]['file_name'] for idx in index_names]}")
                            
                            query_embedding = embedding_service.embed_query(question)
                            search_result = pinecone_service.search_across_indexes(
                                query_embedding=query_embedding,
                                index_names=index_names,
                                top_k=3
                            )
                            
                            if search_result.get("success") and search_result.get("results"):
                                top_results = search_result["results"][:5]
                                
                                best_index = None
                                best_score = 0.0
                                for result in top_results:
                                    if result["score"] > best_score:
                                        best_score = result["score"]
                                        best_index = result["index_name"]
                                
                                best_file_info = file_info_map.get(best_index, {})
                                print(f"[PINECONE] Best match: {best_file_info.get('file_name', 'Unknown')} (score: {best_score:.3f})")
                                
                                chunk_texts = []
                                for result in top_results:
                                    metadata = result.get("metadata", {})
                                    chunk_text = metadata.get("text", "")
                                    file_name = metadata.get("file_name", "Unknown")
                                    chunk_id = result.get("chunk_id", result.get("id", "?"))
                                    score = result.get("score", 0.0)
                                    
                                    chunk_texts.append(f"[Chunk {chunk_id} from {file_name} (score: {score:.3f})]\n{chunk_text}")
                                
                                context_text = "\n\n---\n\n".join(chunk_texts)
                                print(f"[PINECONE] Retrieved {len(top_results)} relevant chunks (best score: {best_score:.3f})")
                                kb_context_found = True
                            else:
                                print(f"[PINECONE] No relevant results found in knowledge base.")
                        else:
                            print(f"[PINECONE] No Pinecone indexes found for knowledge base files.")
                    else:
                        print(f"[PINECONE] Mandatory files referenced by knowledge base are missing or empty.")
                else:
                    print(f"[PINECONE] No files selected in knowledge base.")
            
            except Exception as e:
                print(f"[PINECONE] Error searching Pinecone: {str(e)}")
                import traceback
                print(traceback.format_exc())
                kb_context_found = False
            
            if not kb_context_found:
                use_router_answerer = True
                context_text = ""
        elif use_router_answerer:
            # Use ROUTER + ANSWERER system (Pinecone fallback across all indexes)
            router_data = _search_across_all_files_and_route(question, top_k=10, db=db)
            
            if not router_data['file_scores'] or not router_data['context_chunks']:
                # No results found, fallback to simple response
                error_text = "No relevant documents found in the knowledge base. Please try rephrasing your question or upload relevant files."
                _save_chat_message(db, chat_id, "assistant", error_text, user_email)
                return {
                    "success": False,
                    "error": error_text,
                    "chat_id": chat_id
                }
            
            # Build the prompt for ROUTER + ANSWERER
            user_prompt = f"""Process the following question using the provided file scores and context chunks.

{{
  "user_question": "{question}",
  "file_scores": {json.dumps(router_data['file_scores'], indent=2)},
  "context_chunks": {json.dumps(router_data['context_chunks'], indent=2)}
}}

Return ONLY the JSON object matching the schema specified in your instructions."""
            
            # Send to Gemini with ROUTER + ANSWERER system prompt
            messages = [
                {"role": "system", "content": _router_answerer_system_prompt()},
                {"role": "user", "content": user_prompt}
            ]
            
            print(f"[ROUTER] Sending to LLM with {len(router_data['file_scores'])} files, {len(router_data['context_chunks'])} chunks")
            result = gemini_service.chat(messages, max_tokens=4000)
            
            if result['success']:
                llm_response = result.get('response', '')
                
                # Try to extract JSON from response (may be wrapped in markdown code blocks)
                import re
                json_match = re.search(r'\{[\s\S]*\}', llm_response)
                if json_match:
                    json_str = json_match.group(0)
                else:
                    json_str = llm_response
                
                try:
                    router_result = json.loads(json_str)
                    
                    # Extract answer and format response
                    answer_text = router_result.get('answer', '')
                    status = router_result.get('status', 'OK')
                    
                    # Add sources section if available
                    sources = router_result.get('sources', [])
                    if sources:
                        answer_text += "\n\n**Sources:**\n"
                        for source in sources:
                            answer_text += f"- {source.get('file_name', 'Unknown')} (chunk: {source.get('chunk_id', '?')}, score: {source.get('score', 0):.3f})\n"
                    
                    # Format as HTML for frontend
                    formatted_answer = answer_text.replace('\n', '<br/>')
                    
                    print(f"[ROUTER] LLM response received, status: {status}")
                    print(f"[ROUTER] Selected files: {router_result.get('selected_files', [])}")
                    
                    formatted_answer = formatted_answer or ""
                    _save_chat_message(db, chat_id, "assistant", formatted_answer, user_email)
                    return {
                        "success": True,
                        "response": formatted_answer,
                        "router_result": router_result,  # Include full router result for debugging
                        "status": status,
                        "chat_id": chat_id
                    }
                except json.JSONDecodeError as e:
                    print(f"[ROUTER] Failed to parse JSON response: {e}")
                    print(f"[ROUTER] Raw response: {llm_response[:500]}")
                    # Fallback: return the raw response
                    cleaned_response = llm_response.replace('\n', '<br/>')
                    _save_chat_message(db, chat_id, "assistant", cleaned_response, user_email)
                    return {
                        "success": True,
                        "response": cleaned_response,
                        "status": "OK",
                        "chat_id": chat_id
                    }
            else:
                error_msg = result.get('error', 'Unknown error')
                print(f"[ROUTER] Gemini API error: {error_msg}")
                assistant_error = f"Gemini API error: {error_msg}"
                _save_chat_message(db, chat_id, "assistant", assistant_error, user_email)
                return {
                    "success": False,
                    "error": assistant_error,
                    "chat_id": chat_id
                }
        elif context_text:
            prompt = f"""You are a helpful project management assistant. Based on the following document content, please answer the user's question in a clear, structured, and concise manner.

DOCUMENT CONTENT:
{context_text}

USER QUESTION:
{question}

INSTRUCTIONS:
- Provide a well-structured answer based on the document content
- Use headings, bullet points, and clear formatting
- Focus on answering the user's specific question
- If the question cannot be answered using the document, clearly state that
- Do NOT simply repeat the document content - synthesize and summarize the relevant information
- When multiple documents are provided, consider information from all of them
- **CRITICAL: The document contains links in the format "link_text (url)" or "[Link: url]". You MUST preserve ALL external links from the source document in your response.**
- **When you mention any item that has a link in the source (like "Link to standard documentation/templates", "MOM Template", "RAID Log", "Project Plan", etc.), you MUST include the actual clickable HTML link in the format: <a href="url" target="_blank">link_text</a>**
- **If a link appears in the format "link_text (url)" in the document content where url starts with http:// or https://, convert it to: <a href="url" target="_blank">link_text</a> in your response**
- **If you see "[Link: url]" format in the document, convert it to: <a href="url" target="_blank">View Document</a> or <a href="url" target="_blank">Link</a>**
- **IMPORTANT: If a section mentions a link text (like "Link to sample design document") and you can find a URL in the document content (even if not directly next to the text), include that URL as a clickable link.**
- **Always scan the document content for any URLs (especially Google Sheets links like https://docs.google.com/spreadsheets/...) and include them as clickable links when they relate to the content being discussed.**
- **For links marked as internal (format: "link_text (#internal:...)"), you can skip including the URL but still mention the link text if relevant.**
- **Do NOT skip external links. If a section mentions a link with a valid URL, include that link in your response.**
- **Look for patterns like "Link to...", "...Template", "...Log", "...Plan" - these are likely link references that need to be included with their URLs if they have valid external URLs.**
- **For any Google Sheets or document links (URLs starting with http:// or https://), always include them as clickable links.**
- **Example 1: If you see "Link to standard documentation/templates (https://example.com/templates)" in the document, your response should include: <a href="https://example.com/templates" target="_blank">Link to standard documentation/templates</a>**
- **Example 2: If you see "Link to sample design document" and later find "https://docs.google.com/spreadsheets/d/1Gg4W2tmwaWqFQHpTqFxk3EVdWVuLKFrz/edit..." in the document, include: <a href="https://docs.google.com/spreadsheets/d/1Gg4W2tmwaWqFQHpTqFxk3EVdWVuLKFrz/edit..." target="_blank">Link to sample design document</a>**

Please provide your answer:"""
        else:
            # No context available, answer without document reference
            prompt = question
        
        # Debug: Log prompt length (but not the full content to avoid cluttering logs)
        print(f"[ASK-QUESTION] Prompt length: {len(prompt)} characters")
        print(f"[ASK-QUESTION] Question: {question[:100]}...")
        
        # Send to Gemini
        messages = [
            {"role": "system", "content": _get_structured_html_system_prompt()},
            {"role": "user", "content": prompt}
        ]
        
        result = gemini_service.chat(messages, max_tokens=3000)
        
        if result['success']:
            # Ensure we return the LLM response, not the file context
            llm_response = result.get('response', '')
            
            # Clean up markdown code blocks if present (Gemini sometimes wraps HTML in ```html blocks)
            if llm_response and '```' in llm_response:
                import re
                # Remove opening code fences (```html, ```, etc.)
                llm_response = re.sub(r'```[a-zA-Z]*\s*\n?', '', llm_response)
                # Remove closing code fences
                llm_response = re.sub(r'```\s*\n?', '', llm_response)
                llm_response = llm_response.strip()
                print(f"[ASK-QUESTION] Cleaned markdown code blocks from response")
            
            # Debug: Log response length to ensure we're getting LLM output
            print(f"[ASK-QUESTION] LLM Response length: {len(llm_response)} characters")
            print(f"[ASK-QUESTION] LLM Response preview: {llm_response[:200]}...")
            
            # Ensure response is not empty
            if not llm_response or not llm_response.strip():
                llm_response = "<p>I apologize, but I couldn't generate a response. Please try again.</p>"
                print(f"[ASK-QUESTION] Empty response detected, using fallback message")
            
            _save_chat_message(db, chat_id, "assistant", llm_response, user_email)
            return {
                "success": True,
                "response": llm_response,
                "file_id": file_id,
                "chat_id": chat_id
            }
        else:
            error_msg = result.get('error', 'Unknown error')
            print(f"[ASK-QUESTION] Gemini API error: {error_msg}")
            assistant_error = f"Gemini API error: {error_msg}"
            _save_chat_message(db, chat_id, "assistant", assistant_error, user_email)
            return {
                "success": False,
                "error": assistant_error,
                "response": result.get('response', 'Failed to get response'),
                "chat_id": chat_id
            }
            
    except Exception as e:
        error_message = f"Error processing question: {str(e)}"
        _save_chat_message(db, chat_id, "assistant", error_message, user_email)
        return {
            "success": False,
            "error": error_message,
            "chat_id": chat_id
        }


@app.get("/api/chat/sessions")
async def get_chat_sessions(user_email: str = None, db: Session = Depends(get_db)):
    """Return a list of distinct chat sessions for the current user from conversations table.
    Only returns conversations that don't belong to a project (project_id IS NULL).
    Project conversations are shown under their respective projects."""
    try:
        import json
        
        # Query conversations table - only get conversations WITHOUT project_id
        query = db.query(Conversation).filter(Conversation.project_id.is_(None))
        
        # Only show conversations that belong to the logged-in user
        # Do NOT show conversations with NULL user_email (old chats without user assignment)
        if user_email:
            query = query.filter(Conversation.user_email == user_email)
        else:
            # If no user_email provided, return empty list (require authentication)
            return {
                "success": True,
                "chats": []
            }
        
        conversations = query.order_by(Conversation.updated_at.desc()).all()
        
        sessions = []
        for conv in conversations:
            # Extract preview from first message in conversation JSON
            preview = ""
            try:
                conv_data = conv.conversation_json
                if not isinstance(conv_data, dict):
                    conv_data = json.loads(conv_data) if isinstance(conv_data, str) else {}
                
                messages = conv_data.get("messages", [])
                if messages and len(messages) > 0:
                    first_msg = messages[0]
                    # Get first non-empty user or assistant message
                    preview = first_msg.get("user", "") or first_msg.get("assistant", "")
                    if preview:
                        preview = preview.strip()[:120]
            except Exception as e:
                print(f"[CHAT] Error extracting preview: {str(e)}")
            
            sessions.append({
                "chat_id": conv.chat_id,
                "conversation_id": conv.id,
                "first_message_preview": preview,
                "last_message_at": conv.updated_at.isoformat() if conv.updated_at else None,
                "updated_at": conv.updated_at.isoformat() if conv.updated_at else None
            })
        
        return {
            "success": True,
            "chats": sessions
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error fetching chat sessions: {str(e)}"
        }


@app.get("/api/chat/messages")
async def get_chat_messages(chat_id: str, user_email: str = None, db: Session = Depends(get_db)):
    """Return the ordered messages for a specific chat session from conversations table."""
    if not chat_id:
        return {
            "success": False,
            "error": "chat_id is required"
        }
    
    try:
        import json
        
        # Query conversation by chat_id
        query = db.query(Conversation).filter(Conversation.chat_id == chat_id)
        # Only show conversations that belong to the logged-in user
        # Do NOT show conversations with NULL user_email (old chats without user assignment)
        if user_email:
            query = query.filter(Conversation.user_email == user_email)
        
        conversation = query.first()
        
        if not conversation:
            return {
                "success": True,
                "messages": [],
                "project": None
            }
        
        # Get project info if this conversation belongs to a project
        project_info = None
        if conversation.project_id:
            project = db.query(Project).filter(Project.id == conversation.project_id).first()
            if project:
                project_info = {
                    "id": project.id,
                    "name": project.name
                }
        
        # Extract messages from conversation JSON
        conv_data = conversation.conversation_json
        if not isinstance(conv_data, dict):
            conv_data = json.loads(conv_data) if isinstance(conv_data, str) else {}
        
        messages_json = conv_data.get("messages", [])
        
        # Convert JSON format to message list format
        messages = []
        for msg_obj in messages_json:
            message_id = msg_obj.get("message_id", 0)
            user_msg = msg_obj.get("user", "").strip()
            assistant_msg = msg_obj.get("assistant", "").strip()
            
            # Add user message if exists
            if user_msg:
                messages.append({
                    "id": f"{conversation.id}-{message_id}-user",
                    "chat_id": conversation.chat_id,
                    "user_email": conversation.user_email,
                    "role": "user",
                    "message": user_msg,
                    "created_at": conversation.created_at.isoformat() if conversation.created_at else None
                })
            
            # Add assistant message if exists
            if assistant_msg:
                messages.append({
                    "id": f"{conversation.id}-{message_id}-assistant",
                    "chat_id": conversation.chat_id,
                    "user_email": conversation.user_email,
                    "role": "assistant",
                    "message": assistant_msg,
                    "created_at": conversation.updated_at.isoformat() if conversation.updated_at else None
                })
        
        return {
            "success": True,
            "messages": messages,
            "project": project_info
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error fetching chat messages: {str(e)}"
        }


@app.post("/api/chat/save-message")
async def save_chat_message(
    chat_id: str = Form(...),
    role: str = Form(...),
    message: str = Form(...),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """Save a single chat message in real-time by updating conversation JSON."""
    try:
        import json
        
        # Find or create conversation record
        conversation = db.query(Conversation).filter(
            Conversation.chat_id == chat_id
        ).first()
        
        if conversation:
            # Update existing conversation
            conversation_data = conversation.conversation_json
            if not isinstance(conversation_data, dict):
                conversation_data = json.loads(conversation_data) if isinstance(conversation_data, str) else {}
            
            messages = conversation_data.get("messages", [])
            
            # Find the last message entry or create new one
            if messages and role == "assistant" and len(messages) > 0:
                # If assistant message, update the last message entry
                last_msg = messages[-1]
                if "assistant" not in last_msg or not last_msg.get("assistant"):
                    last_msg["assistant"] = message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": "",
                        "assistant": message
                    })
            elif messages and role == "user":
                # If user message, check if last message has user or create new
                last_msg = messages[-1] if messages else None
                if last_msg and not last_msg.get("user"):
                    last_msg["user"] = message
                else:
                    # Create new message entry
                    new_message_id = len(messages) + 1
                    messages.append({
                        "message_id": new_message_id,
                        "user": message,
                        "assistant": ""
                    })
            else:
                # First message
                messages.append({
                    "message_id": 1,
                    "user": message if role == "user" else "",
                    "assistant": message if role == "assistant" else ""
                })
            
            conversation_data["messages"] = messages
            conversation.conversation_json = conversation_data
            if user_email:
                conversation.user_email = user_email
            db.commit()
        else:
            # Create new conversation
            conversation_id = str(uuid.uuid4())
            conversation_data = {
                "conversation_id": 1,  # Will be updated to actual id after save
                "messages": [{
                    "message_id": 1,
                    "user": message if role == "user" else "",
                    "assistant": message if role == "assistant" else ""
                }]
            }
            new_conversation = Conversation(
                conversation_id=conversation_id,
                chat_id=chat_id,
                user_email=user_email,
                conversation_json=conversation_data
            )
            db.add(new_conversation)
            db.commit()
            db.refresh(new_conversation)
            # Update conversation_id in JSON to actual database id
            conversation_data["conversation_id"] = new_conversation.id
            new_conversation.conversation_json = conversation_data
            db.commit()
        
        return {
            "success": True,
            "message": "Message saved successfully"
        }
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error saving message: {str(e)}"
        }


@app.post("/api/chat/save-conversation")
async def save_conversation_json(
    chat_id: str = Form(...),
    conversation_json: str = Form(...),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """Save full conversation as JSON for easy retrieval."""
    try:
        import json
        conversation_data = json.loads(conversation_json)
        
        # Find or create conversation record
        conversation = db.query(Conversation).filter(
            Conversation.chat_id == chat_id
        ).first()
        
        if conversation:
            # Update existing conversation
            # Ensure conversation_id is set correctly
            if "conversation_id" not in conversation_data or conversation_data["conversation_id"] != conversation.id:
                conversation_data["conversation_id"] = conversation.id
            
            conversation.conversation_json = conversation_data
            if user_email:
                conversation.user_email = user_email
            db.commit()
        else:
            # Create new conversation
            conversation_id = str(uuid.uuid4())
            # Ensure conversation_id is set in the JSON
            if "conversation_id" not in conversation_data:
                conversation_data["conversation_id"] = 1  # Will be updated after save
            
            new_conversation = Conversation(
                conversation_id=conversation_id,
                chat_id=chat_id,
                user_email=user_email,
                conversation_json=conversation_data
            )
            db.add(new_conversation)
            db.commit()
            db.refresh(new_conversation)
            # Update conversation_id in JSON to actual database id
            conversation_data["conversation_id"] = new_conversation.id
            new_conversation.conversation_json = conversation_data
            db.commit()
        
        return {
            "success": True,
            "message": "Conversation saved successfully"
        }
    except Exception as e:
        db.rollback()
        return {
            "success": False,
            "error": f"Error saving conversation: {str(e)}"
        }


@app.post("/api/chat/create")
async def create_chat_record(
    chat_id: str = Form(...),
    user_email: str = Form(None),
    db: Session = Depends(get_db)
):
    """Create a new chat record when a chat session starts."""
    try:
        # Create an initial message to establish the chat record
        # The chat_id will be used to group messages
        print(f"[CHAT] Creating chat record for chat_id: {chat_id}, user: {user_email}")
        
        # The chat record is automatically created when the first message is saved
        # This endpoint just ensures the chat_id exists in the system
        return {
            "success": True,
            "message": "Chat record created successfully",
            "chat_id": chat_id
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Error creating chat record: {str(e)}"
        }

# Feedback endpoints
@app.post("/api/feedback", response_model=FeedbackResponse)
async def submit_feedback(feedback_data: FeedbackRequest, db: Session = Depends(get_db)):
    """Submit feedback form data"""
    try:
        print(f"[FEEDBACK] Received feedback submission")
        print(f"[FEEDBACK] Data: {feedback_data.dict()}")
        
        # Create new feedback record
        feedback = Feedback(
            name=feedback_data.name,
            email=feedback_data.email,
            additional_comments=feedback_data.additional_comments,
            created_by=feedback_data.user_email or feedback_data.email
        )
        
        # Add to database
        db.add(feedback)
        db.commit()
        db.refresh(feedback)
        
        print(f"[FEEDBACK] Successfully saved feedback with ID: {feedback.id}")
        
        return FeedbackResponse(
            success=True,
            message="Feedback submitted successfully!",
            feedback_id=feedback.id
        )
        
    except Exception as e:
        print(f"[FEEDBACK] Error submitting feedback: {str(e)}")
        db.rollback()
        return FeedbackResponse(
            success=False,
            message=f"Error submitting feedback: {str(e)}"
        )

@app.get("/api/feedback")
async def get_feedback(user_email: str = None, db: Session = Depends(get_db)):
    """Get feedback submissions (admin or user-specific)"""
    try:
        if user_email:
            # Get feedback for specific user
            feedback = db.query(Feedback).filter(Feedback.created_by == user_email).all()
        else:
            # Get all feedback (admin)
            feedback = db.query(Feedback).all()
        
        return {
            "success": True,
            "feedback": [
                {
                    "id": f.id,
                    "name": f.name,
                    "email": f.email,
                    "additional_comments": f.additional_comments,
                    "created_at": f.created_at,
                    "created_by": f.created_by
                }
                for f in feedback
            ]
        }
    except Exception as e:
        return {"success": False, "message": str(e)}

# Workspace endpoints
@app.get("/api/workspaces")
async def get_workspaces(db: Session = Depends(get_db)):
    """Get all workspaces"""
    try:
        from models import Workspace
        
        workspaces = db.query(Workspace).order_by(Workspace.is_default.desc(), Workspace.name).all()
        
        return {
            "success": True,
            "workspaces": [
                {
                    "id": ws.id,
                    "name": ws.name,
                    "description": ws.description,
                    "is_default": ws.is_default,
                    "created_at": ws.created_at.isoformat() if ws.created_at else None
                }
                for ws in workspaces
            ]
        }
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.post("/api/workspaces")
async def create_workspace(
    name: str = Form(...),
    description: str = Form(None),
    db: Session = Depends(get_db)
):
    """Create a new workspace"""
    try:
        from models import Workspace
        
        # Check if workspace with same name already exists
        existing_workspace = db.query(Workspace).filter(Workspace.name == name).first()
        if existing_workspace:
            return {"success": False, "message": f"Workspace '{name}' already exists"}
        
        # Create new workspace
        workspace = Workspace(
            name=name,
            description=description,
            is_default=False
        )
        
        db.add(workspace)
        db.commit()
        db.refresh(workspace)
        
        return {
            "success": True,
            "workspace": {
                "id": workspace.id,
                "name": workspace.name,
                "description": workspace.description,
                "is_default": workspace.is_default,
                "created_at": workspace.created_at.isoformat() if workspace.created_at else None
            }
        }
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

@app.delete("/api/workspaces/{workspace_id}")
async def delete_workspace(
    workspace_id: int, 
    user_email: str = None,
    db: Session = Depends(get_db)
):
    """Delete a workspace - Admin only"""
    try:
        from models import Workspace
        
        # Define admin emails
        ADMIN_EMAILS = [
            "shaik.sharuk@forsysinc.com"
            # Add more admin emails here as needed
        ]
        
        # Check if user is admin
        if not user_email or user_email not in ADMIN_EMAILS:
            return {"success": False, "message": "Only admins can delete workspaces"}
        
        # Find the workspace
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        
        if not workspace:
            return {"success": False, "message": "Workspace not found"}
        
        # Prevent deletion of default workspace
        if workspace.is_default:
            return {"success": False, "message": "Cannot delete the default workspace"}
        
        # Delete the workspace
        db.delete(workspace)
        db.commit()
        
        return {"success": True, "message": f"Workspace '{workspace.name}' deleted successfully"}
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

# Projects endpoints
@app.post("/api/projects")
async def create_project(
    name: str = Form(...),
    user_email: str = Form(...),
    db: Session = Depends(get_db)
):
    """Create a new project with a default conversation"""
    try:
        # Validate project name
        if not name or not name.strip():
            return {"success": False, "message": "Project name cannot be empty"}
        
        # Begin transaction - create project and conversation atomically
        # Create new project with unique ID
        project = Project(
            name=name.strip(),
            user_email=user_email
        )
        
        db.add(project)
        db.flush()  # Flush to get project.id without committing
        
        # Create default conversation linked to the project
        conversation_id = str(uuid.uuid4())
        chat_id = str(uuid.uuid4())
        conversation_data = {
            "conversation_id": 0,  # Will be updated to actual id after save
            "messages": []
        }
        
        conversation = Conversation(
            conversation_id=conversation_id,
            chat_id=chat_id,
            user_email=user_email,
            project_id=project.id,
            conversation_json=conversation_data
        )
        
        db.add(conversation)
        db.commit()  # Commit both project and conversation together
        db.refresh(project)
        db.refresh(conversation)
        
        # Update conversation_id in JSON to actual database id
        conversation_data["conversation_id"] = conversation.id
        conversation.conversation_json = conversation_data
        db.commit()
        db.refresh(conversation)
        
        return {
            "success": True,
            "project": {
                "id": project.id,
                "name": project.name,
                "user_email": project.user_email,
                "created_at": project.created_at.isoformat() if project.created_at else None
            },
            "conversation": {
                "id": conversation.id,
                "conversation_id": conversation.conversation_id,
                "chat_id": conversation.chat_id,
                "title": "Default chat",
                "project_id": conversation.project_id,
                "user_email": conversation.user_email,
                "created_at": conversation.created_at.isoformat() if conversation.created_at else None
            }
        }
    except Exception as e:
        db.rollback()
        error_msg = str(e)
        import traceback
        error_trace = traceback.format_exc()
        print(f"[PROJECT CREATE ERROR] {error_msg}")
        print(f"[PROJECT CREATE TRACEBACK]\n{error_trace}")
        
        # Check if it's the UUID/integer type error
        if "invalid input syntax for type integer" in error_msg or "InvalidTextRepresentation" in error_msg:
            error_msg = "Database schema error: Projects table has wrong ID type. Please restart the backend server to fix this automatically."
        elif "foreign key" in error_msg.lower() or "constraint" in error_msg.lower():
            error_msg = f"Database constraint error: {error_msg}. Please check foreign key relationships."
        elif "conversations" in error_msg.lower() and "project_id" in error_msg.lower():
            error_msg = f"Conversation creation error: {error_msg}. Please check conversations table schema."
        
        return {"success": False, "message": error_msg}

@app.get("/api/projects")
async def get_projects(
    user_email: str,
    db: Session = Depends(get_db)
):
    """Get all projects for a user with their conversations"""
    try:
        projects = db.query(Project).filter(Project.user_email == user_email).order_by(Project.created_at.desc()).all()
        
        projects_list = []
        for project in projects:
            # Get conversations for this project
            conversations = db.query(Conversation).filter(
                Conversation.project_id == project.id
            ).order_by(Conversation.created_at.asc()).all()
            
            projects_list.append({
                "id": project.id,
                "name": project.name,
                "user_email": project.user_email,
                "created_at": project.created_at.isoformat() if project.created_at else None,
                "conversations": [
                    {
                        "id": conv.id,
                        "conversation_id": conv.conversation_id,
                        "chat_id": conv.chat_id,
                        "title": "Default chat",  # Could be extracted from conversation_json if stored
                        "project_id": conv.project_id,
                        "created_at": conv.created_at.isoformat() if conv.created_at else None
                    }
                    for conv in conversations
                ]
            })
        
        return {
            "success": True,
            "projects": projects_list
        }
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.delete("/api/projects/{project_id}")
async def delete_project(
    project_id: str,
    user_email: str,
    db: Session = Depends(get_db)
):
    """Delete a project"""
    try:
        project = db.query(Project).filter(Project.id == project_id, Project.user_email == user_email).first()
        
        if not project:
            return {"success": False, "message": "Project not found or you don't have permission to delete it"}
        
        db.delete(project)
        db.commit()
        
        return {"success": True, "message": "Project deleted successfully"}
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

@app.get("/api/workspaces/default")
async def get_default_workspace(db: Session = Depends(get_db)):
    """Get the default EJM workspace"""
    try:
        from models import Workspace
        
        # Check if EJM workspace exists
        workspace = db.query(Workspace).filter(Workspace.name == "EJM").first()
        
        if not workspace:
            # Create EJM workspace if it doesn't exist
            workspace = Workspace(
                name="EJM",
                description="Default EJM workspace",
                is_default=True
            )
            db.add(workspace)
            db.commit()
            db.refresh(workspace)
        
            return {
                "success": True, 
            "workspace": {
                "id": workspace.id,
                "name": workspace.name,
                "description": workspace.description,
                "is_default": workspace.is_default
            }
        }
    except Exception as e:
        return {"success": False, "message": str(e)}


if __name__ == "__main__":
    import uvicorn
    # Use 0.0.0.0 to allow external connections on the same network
    uvicorn.run(app, host="0.0.0.0", port=8000)

async def generate_risk_assessment_pdf(risk_assessment_content: str, assessment_name: str):
    """Generate PDF from risk assessment content using ReportLab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from io import BytesIO
        from datetime import datetime
        import re
        
        print("Starting PDF generation for risk assessment with ReportLab...")
        
        # Clean up the HTML content
        cleaned_content = risk_assessment_content.strip()
        
        # Remove markdown code blocks
        if cleaned_content.startswith('```html'):
            cleaned_content = cleaned_content.replace('```html', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        elif cleaned_content.startswith('```'):
            cleaned_content = cleaned_content.replace('```', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        
        # Extract body content if full HTML
        body_match = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            cleaned_content = body_match.group(1)
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        # Container for PDF elements
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Title style with risk assessment colors
        title_style = ParagraphStyle(
            'RiskAssessmentTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#e53e3e'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Add title
        story.append(Paragraph(assessment_name, title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", 
                              styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Parse HTML content properly to handle tables
        story.extend(await parse_html_content_for_pdf(cleaned_content, styles))
        
        # Add footer
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("This risk assessment was generated and shared via PM Portal", footer_style))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        print(f"PDF generated successfully using ReportLab, size: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        print(f"ERROR in PDF generation for risk assessment: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

async def generate_html_email_content(sprint_plan_content: str, sprint_plan_name: str, custom_message: str = ""):
    """Generate HTML email content with proper table structure and formatting"""
    try:
        import re
        from datetime import datetime
        
        print("Generating HTML email content...")
        
        # Clean up the HTML content
        cleaned_content = sprint_plan_content.strip()
        
        # Remove markdown code blocks
        if cleaned_content.startswith('```html'):
            cleaned_content = cleaned_content.replace('```html', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        elif cleaned_content.startswith('```'):
            cleaned_content = cleaned_content.replace('```', '', 1)
            cleaned_content = cleaned_content.rsplit('```', 1)[0]
        
        # Extract body content if full HTML
        body_match = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            cleaned_content = body_match.group(1)
        
        # Wrap all tables in table-wrapper divs for better responsive handling
        cleaned_content = re.sub(
            r'<table([^>]*)>(.*?)</table>',
            r'<div class="table-wrapper"><table\1>\2</table></div>',
            cleaned_content,
            flags=re.DOTALL | re.IGNORECASE
        )
        
        # Create HTML email template
        html_email = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Sprint Plan - {sprint_plan_name}</title>
            <style>
                body {{
                    font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, sans-serif;
                    margin: 0;
                    padding: 10px;
                    background: #f8fafc;
                    color: #2c3e50;
                    line-height: 1.6;
                    font-size: 14px;
                    overflow-x: auto;
                }}
                
                .email-container {{
                    max-width: 100%;
                    margin: 0 auto;
                    background: white;
                    border-radius: 12px;
                    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
                    overflow-x: auto;
                }}
                
                .email-header {{
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 30px;
                    text-align: center;
                }}
                
                .email-header h1 {{
                    margin: 0 0 10px 0;
                    font-size: 2em;
                    font-weight: 700;
                }}
                
                .email-header p {{
                    margin: 0;
                    opacity: 0.9;
                    font-size: 1.1em;
                }}
                
                .email-body {{
                    padding: 30px;
                }}
                
                .custom-message {{
                    background: #e6fffa;
                    border: 1px solid #81e6d9;
                    border-radius: 8px;
                    padding: 20px;
                    margin-bottom: 30px;
                    color: #234e52;
                }}
                
                .custom-message h3 {{
                    margin: 0 0 10px 0;
                    color: #2d3748;
                    font-size: 1.2em;
                }}
                
                .sprint-content {{
                    background: #f8fafc;
                    border: 1px solid #e2e8f0;
                    border-radius: 8px;
                    padding: 25px;
                }}
                
                .sprint-content h2, .sprint-content h3, .sprint-content h4 {{
                    color: #2d3748;
                    margin: 20px 0 15px 0;
                    font-weight: 600;
                }}
                
                .sprint-content h2 {{
                    font-size: 1.4em;
                    color: #1a202c;
                    border-bottom: 2px solid #e2e8f0;
                    padding-bottom: 10px;
                }}
                
                .sprint-content h3 {{
                    font-size: 1.2em;
                    color: #2d3748;
                }}
                
                .sprint-content h4 {{
                    font-size: 1.1em;
                    color: #4a5568;
                }}
                
                .sprint-content p {{
                    margin: 12px 0;
                    line-height: 1.6;
                    color: #4a5568;
                }}
                
                .sprint-content ul, .sprint-content ol {{
                    margin: 15px 0;
                    padding-left: 25px;
                }}
                
                .sprint-content li {{
                    margin: 8px 0;
                    line-height: 1.6;
                    color: #4a5568;
                }}
                
                .sprint-content strong {{
                    color: #2d3748;
                    font-weight: 600;
                }}
                
                .sprint-content em {{
                    color: #718096;
                    font-style: italic;
                }}
                
                .sprint-content table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    background: white;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    table-layout: auto;
                    min-width: 100%;
                }}
                
                .sprint-content th {{
                    background: #667eea;
                    color: white;
                    padding: 15px 12px;
                    text-align: left;
                    font-weight: 600;
                    font-size: 0.9em;
                    text-transform: uppercase;
                    letter-spacing: 0.5px;
                    white-space: normal;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    vertical-align: top;
                    min-width: 120px;
                }}
                
                .sprint-content td {{
                    padding: 15px 12px;
                    border-bottom: 1px solid #e2e8f0;
                    color: #4a5568;
                    white-space: normal;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    font-size: 0.9em;
                    line-height: 1.5;
                    vertical-align: top;
                    min-width: 120px;
                }}
                
                .sprint-content tr:nth-child(even) {{
                    background: #f8fafc;
                }}
                
                .sprint-content tr:hover {{
                    background: #edf2f7;
                }}
                
                /* Responsive table wrapper */
                .table-wrapper {{
                    overflow-x: auto;
                    margin: 20px 0;
                    border-radius: 8px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    width: 100%;
                }}
                
                .table-wrapper table {{
                    margin: 0;
                    min-width: 800px;
                    width: 100%;
                }}
                
                /* Specific column width improvements */
                .sprint-content table td:nth-child(1),
                .sprint-content table th:nth-child(1) {{
                    min-width: 100px;
                    max-width: 150px;
                }}
                
                .sprint-content table td:nth-child(2),
                .sprint-content table th:nth-child(2) {{
                    min-width: 200px;
                    max-width: 300px;
                }}
                
                .sprint-content table td:nth-child(3),
                .sprint-content table th:nth-child(3) {{
                    min-width: 120px;
                    max-width: 180px;
                }}
                
                .sprint-content table td:nth-child(4),
                .sprint-content table th:nth-child(4) {{
                    min-width: 250px;
                    max-width: 400px;
                }}
                
                .sprint-content table td:nth-child(5),
                .sprint-content table th:nth-child(5) {{
                    min-width: 120px;
                    max-width: 180px;
                }}
                
                .email-footer {{
                    background: #f7fafc;
                    padding: 20px 30px;
                    text-align: center;
                    color: #718096;
                    font-size: 0.9em;
                    border-top: 1px solid #e2e8f0;
                }}
                
                .attachment-notice {{
                    background: #fff5f5;
                    border: 1px solid #feb2b2;
                    border-radius: 8px;
                    padding: 15px;
                    margin-top: 20px;
                    color: #742a2a;
                }}
                
                .attachment-notice strong {{
                    color: #c53030;
                }}
                
                
                /* Scroll hint */
                .scroll-hint {{
                    background: #e6fffa;
                    border: 1px solid #81e6d9;
                    border-radius: 6px;
                    padding: 10px;
                    margin: 10px 0;
                    color: #234e52;
                    font-size: 0.9em;
                    text-align: center;
                }}
            </style>
        </head>
        <body>
            <div class="email-container">
                <div class="email-header">
                    <h1>Sprint Plan - {sprint_plan_name}</h1>
                    <p>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
                </div>
                
                <div class="email-body">
                    {f'<div class="custom-message"><h3>Message from sender:</h3><p>{custom_message}</p></div>' if custom_message else ''}
                    
                    <div class="scroll-hint">
                        <strong>Tip:</strong> If tables appear cut off, scroll horizontally to see all columns
                    </div>
                    
                    <div class="sprint-content">
                        {cleaned_content}
                    </div>
                    
                    <div class="attachment-notice">
                        <strong>PDF Attachment:</strong> A formatted PDF version of this sprint plan is attached to this email for your convenience.
                    </div>
                </div>
                
                <div class="email-footer">
                    <p>This sprint plan was generated and shared via PM Portal</p>
                    <p>For any questions or concerns, please contact the sender.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        print("HTML email content generated successfully")
        return html_email
        
    except Exception as e:
        print(f"ERROR generating HTML email content: {str(e)}")
        import traceback
        traceback.print_exc()
        # Fallback to simple HTML
        return f"""
        <html>
        <body>
            <h2>Sprint Plan - {sprint_plan_name}</h2>
            <p>{custom_message}</p>
            <div>{sprint_plan_content}</div>
            <p><strong>Note:</strong> A PDF attachment is also included with this email.</p>
        </body>
        </html>
        """

async def parse_html_content_for_pdf(html_content: str, styles):
    """Parse HTML content and convert to ReportLab elements, properly handling tables"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from io import BytesIO
        from datetime import datetime
        import re
        
        elements = []
        
        # Clean up HTML
        cleaned_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)
        cleaned_content = re.sub(r'<script[^>]*>.*?</script>', '', cleaned_content, flags=re.DOTALL)
        
        # Split content by tables and other elements
        parts = re.split(r'(<table[^>]*>.*?</table>)', cleaned_content, flags=re.DOTALL | re.IGNORECASE)
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            # Check if this is a table
            if re.match(r'<table[^>]*>', part, re.IGNORECASE):
                # Parse table
                table_elements = parse_html_table_for_pdf(part, styles)
                elements.extend(table_elements)
            else:
                # Parse other content (headings, paragraphs, etc.)
                other_elements = parse_html_text_for_pdf(part, styles)
                elements.extend(other_elements)
        
        return elements
        
    except Exception as e:
        print(f"ERROR parsing HTML content for PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def parse_html_table_for_pdf(table_html: str, styles):
    """Parse HTML table and convert to ReportLab Table"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import re
        
        elements = []
        
        # Extract table rows
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
        
        if not rows:
            return elements
        
        table_data = []
        
        for row in rows:
            # Extract cells (both th and td)
            cells = re.findall(r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>', row, re.DOTALL | re.IGNORECASE)
            row_data = []
            
            for cell in cells:
                # Clean cell content
                cell_text = re.sub(r'<[^>]+>', '', cell)  # Remove HTML tags
                cell_text = cell_text.strip()
                if not cell_text:
                    cell_text = " "
                
                # Handle long text better - split into multiple lines if needed
                if len(cell_text) > 80:
                    # Split long text into multiple lines for better readability
                    words = cell_text.split()
                    lines = []
                    current_line = ""
                    
                    for word in words:
                        if len(current_line + " " + word) <= 80:
                            current_line += (" " + word) if current_line else word
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        lines.append(current_line)
                    
                    # Join lines with line breaks for PDF
                    cell_text = "\n".join(lines)
                
                row_data.append(cell_text)
            
            if row_data:
                table_data.append(row_data)
        
        if table_data:
            # Calculate column widths based on content and page width
            page_width = A4[0] - 144  # A4 width minus margins (72*2)
            num_cols = len(table_data[0]) if table_data else 1
            
            # Calculate column widths based on number of columns
            if num_cols <= 2:
                # For 1-2 columns, use wider columns
                col_widths = [page_width / num_cols] * num_cols
            elif num_cols <= 4:
                # For 3-4 columns, use moderate width
                col_widths = [page_width / num_cols] * num_cols
            else:
                # For 5+ columns, use smaller width but ensure minimum readability
                max_col_width = page_width / num_cols
                min_col_width = 60  # Minimum column width for readability
                if max_col_width < min_col_width:
                    # If columns would be too narrow, use minimum width and allow horizontal overflow
                    col_widths = [min_col_width] * num_cols
                else:
                    col_widths = [max_col_width] * num_cols
            
            # Create ReportLab Table with calculated column widths
            pdf_table = Table(table_data, colWidths=col_widths)
            
            # Style the table
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),  # Header background
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # Header text color
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Left align all cells
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Bold header
                ('FONTSIZE', (0, 0), (-1, 0), 9),  # Header font size
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),  # Regular font for data
                ('FONTSIZE', (0, 1), (-1, -1), 8),  # Data font size
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header padding
                ('TOPPADDING', (0, 0), (-1, 0), 12),  # Header padding
                ('BOTTOMPADDING', (0, 1), (-1, -1), 10),  # Data padding
                ('TOPPADDING', (0, 1), (-1, -1), 10),  # Data padding
                ('LEFTPADDING', (0, 0), (-1, -1), 8),  # Left padding
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),  # Right padding
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),  # Grid lines
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Top align all cells
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),  # Alternating row colors
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#667eea')),  # Header bottom line
                ('LINEBELOW', (0, 1), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),  # Row separator lines
                ('MINROWHEIGHT', (0, 0), (-1, -1), 20),  # Minimum row height
            ])
            
            pdf_table.setStyle(table_style)
            
            # Add spacing before and after table
            elements.append(Spacer(1, 0.15*inch))
            elements.append(pdf_table)
            elements.append(Spacer(1, 0.15*inch))
        
        return elements
        
    except Exception as e:
        print(f"ERROR parsing HTML table for PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return []



async def send_email_with_pdf_via_smtp(to_email: str, subject: str, body: str, from_email: str, pdf_content: bytes, sprint_plan_name: str):
    """Fallback SMTP email sending with PDF attachment"""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.application import MIMEApplication
        import os
        from datetime import datetime
        
        # SMTP configuration
        SMTP_SERVER = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        SMTP_PORT = int(os.getenv('SMTP_PORT', '587'))
        SMTP_USERNAME = os.getenv('SMTP_USERNAME')
        SMTP_PASSWORD = os.getenv('SMTP_PASSWORD')
        
        if not SMTP_USERNAME or not SMTP_PASSWORD:
            return {
                "success": False, 
                "message": "Email configuration not set. Please configure SMTP_USERNAME and SMTP_PASSWORD in environment variables."
            }
        
        # Create message with attachment
        msg = MIMEMultipart()
        msg['Subject'] = subject
        msg['From'] = from_email
        msg['To'] = to_email
        
        # Add body (HTML content)
        msg.attach(MIMEText(body, 'html'))
        
        # Add PDF attachment
        if pdf_content:
            pdf_filename = f"{sprint_plan_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            print(f"Attaching PDF: {pdf_filename}, size: {len(pdf_content)} bytes")
            pdf_attachment = MIMEApplication(pdf_content, _subtype='pdf')
            pdf_attachment.add_header('Content-Disposition', 'attachment', filename=pdf_filename)
            msg.attach(pdf_attachment)
            print("PDF attachment added successfully")
        else:
            print("WARNING: No PDF content to attach!")
        
        # Send email
        print(f"Connecting to SMTP server: {SMTP_SERVER}:{SMTP_PORT}")
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        print("Logging in to SMTP...")
        server.login(SMTP_USERNAME, SMTP_PASSWORD)
        print("Sending email...")
        server.send_message(msg)
        server.quit()
        print("Email sent successfully!")
        
        return {
            "success": True, 
            "message": f"Sprint plan sent successfully to {to_email} with PDF attachment"
        }
        
    except Exception as e:
        return {"success": False, "message": f"SMTP Error: {str(e)}"}

@app.get("/api/workspaces/default")
async def get_default_workspace(db: Session = Depends(get_db)):
    """Get the default EJM workspace"""
    try:
        from models import Workspace
        
        # Check if EJM workspace exists
        workspace = db.query(Workspace).filter(Workspace.name == "EJM").first()
        
        if not workspace:
            # Create EJM workspace if it doesn't exist
            workspace = Workspace(
                name="EJM",
                description="Default EJM workspace",
                is_default=True
            )
            db.add(workspace)
            db.commit()
            db.refresh(workspace)
        
        return {
            "success": True,
            "workspace": {
                "id": workspace.id,
                "name": workspace.name,
                "description": workspace.description,
                "is_default": workspace.is_default
            }
        }
    except Exception as e:
        return {"success": False, "message": str(e)}


if __name__ == "__main__":
    import uvicorn
    # Use 0.0.0.0 to allow external connections on the same network
    uvicorn.run(app, host="0.0.0.0", port=8000)
